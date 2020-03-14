import os
import time
import json

from epyk.core.js import Imports
from epyk.core.js import Js
from epyk.core.js import JsUtils
from epyk.core.js.Imports import requires

from epyk.core.html.templates import HtmlTmplBase


class OutBrowsers(object):
  def __init__(self, context):
    self._context = context

  def codepen(self, path=None, target="_blank", open_browser=True):
    """
    Description:
    ------------
    Update the Html launcher and send the data to codepen
    URL used: https://codepen.io/pen/define/

    Related Pages:
    --------------
    https://www.debuggex.com/cheatsheet/regex/python

    Attributes:
    ----------
    :param path: Output path in which the static files will be generated
    :param target: String. Load the data in a new tab in the browser
    :param open_browser: Boolean. Flag to open the browser automatically

    :return: The output launcher full file name
    """
    import re
    import webbrowser

    results = self._context._to_html_obj(content_only=True)
    js_external = re.findall('<script language="javascript" type="text/javascript" src="(.*?)"></script>', results['jsImports'])
    css_external = re.findall('<link rel="stylesheet" href="(.*?)" type="text/css">', results['cssImports'])
    jsObj = Js.JsBase()
    result = {"js": results["jsFrgs"], "js_external": ";".join(js_external), "css_external": ";".join(css_external), "html": results['content'], "css": results["cssStyle"]}
    data = jsObj.location.postTo("https://codepen.io/pen/define/", {"data": json.dumps(result)}, target=target)
    if path is None:
      path = os.path.join(os.getcwd(), "outs")
    else:
      path = os.path.join(path)
    if not os.path.exists(path):
      os.makedirs(path, exist_ok=True)
    with open(os.path.join(path, "RunnerCodepen.html"), "w") as f:
      f.write('<html><body></body><script>%s</script></html>' % data.replace("\\\\n", ""))
    launcher_file = os.path.join(path, "RunnerCodepen.html")
    if open_browser:
      webbrowser.open(launcher_file)
    return launcher_file


class PyOuts(object):
  def __init__(self, report=None):
    self._report = report
    self.excluded_packages = None

  def _to_html_obj(self, htmlParts=None, cssParts=None):
    """
    Description:
    ------------
    Create the HTML result object from the report definition

    Attributes:
    ----------
    :param htmlParts: Optional. HTML Content of the page
    :param cssParts: Optional. CSS classes content of the page

    :return: A python dictionary with the HTML results
    """
    if htmlParts is None:
      htmlParts, cssParts = [], {}
      for objId in self._report.content:
        #if content_only and self._report.htmlItems[objId].name == "Nav Bar":
        #  continue

        if self._report.htmlItems[objId].inReport:
          htmlParts.append(self._report.htmlItems[objId].html())
        #
        cssParts.update(self._report.htmlItems[objId].style.get_classes_css())
    onloadParts = []
    for data_id, data in self._report._props.get("data", {}).get('sources', {}).items():
      onloadParts.append("var data_%s = %s" % (data_id, json.dumps(data)))

    for k, v in self._report._props.get('js', {}).get('functions', {}).items():
      sPmt = "(%s)" % ", ".join(list(v["pmt"])) if "pmt" in v else "{}"
      onloadParts.append("function %s%s{%s}" % (k, sPmt, v["content"].strip()))

    for c, d in self._report._props.get('js', {}).get("constructors", {}).items():
      onloadParts.append(d)

    for c, d in self._report._props.get('js', {}).get("datasets", {}).items():
      onloadParts.append(d)

    for b in self._report._props.get('js', {}).get("builders", []):
      onloadParts.append(b)

    # Add the component on ready functions
    for objId in self._report.content:
      obj_id = self._report.htmlItems[objId].dom.varId
      if obj_id and obj_id in self._report._props.get('js', {}).get('onCompReady', {}):
        onloadParts.append(self._report._props['js']['onCompReady'][obj_id])
      for event, fncs in self._report.htmlItems[objId]._events['doc_ready'].items():
        if fncs['profile']:
          # Add the profiling feature if the variable is set to true
          profile_var = self._report.js.performance.add_profiling(fncs['content'])
          fncs['content'].append(self._report.js.console.log(self._report.js.objects.get(profile_var)))
        str_fncs = JsUtils.jsConvertFncs(fncs['content'], toStr=True)
        onloadParts.append("%s.addEventListener('%s', function(event){%s})" % (obj_id, event, str_fncs))

    # Add the page on document ready functions
    for on_ready_frg in self._report._props.get('js', {}).get('onReady', []):
      onloadParts.append(on_ready_frg)

    importMng = Imports.ImportManager(online=True, report=self._report)
    results = {
      'cssStyle': "%s\n%s" % ("\n".join([v for v in cssParts.values()]), self._report._cssText),
      'cssContainer': ";".join(["%s:%s" % (k, v) for k, v in self._report._props.get('css', {}).get('container', {}).items()]),
      'content': "\n".join(htmlParts),
      'jsFrgs': ";".join(onloadParts),
      'cssImports': importMng.cssResolve(self._report.cssImport, self._report.cssLocalImports, excluded=self.excluded_packages),
      'jsImports': importMng.jsResolve(self._report.jsImports, self._report.jsLocalImports, excluded=self.excluded_packages)
    }
    return results

  def _repr_html_(self):
    """
    Description:
    ------------
    Standard output for Jupyter Notebooks.

    This is what will use IPython in order to display the results in cells.
    """
    results = self._to_html_obj(content_only=True)
    return HtmlTmplBase.JUPYTER.strip() % results

  def jupyterlab(self):
    """
    Description:
    ------------
    For a display of the report in JupyterLab.
    Thanks to this function some packages will not be imported to not conflict with the existing ones

    Related Pages:
    --------------
    https://jupyter.org/
    """
    self.excluded_packages = ['bootstrap']
    return self

  def jupyter(self):
    """
    Description:
    ------------
    For a display of the report in Jupyter.
    Thanks to this function some packages will not be imported to not conflict with the existing ones

    Related Pages:
    --------------
    https://jupyter.org/

    :return: The ouput object with the function _repr_html_
    """
    self.excluded_packages = ['bootstrap', 'jquery']
    return self

  def w3cTryIt(self, path=None, name=None):
    """
    Description:
    ------------
    This will produce everything in a single page which can be directly copied to the try editor in w3C website

    Related Pages:
    --------------
    https://www.w3schools.com/html/tryit.asp?filename=tryhtml_basic

    Attributes:
    ----------
    :param path: The path in which the output files will be created
    :param name: The filename without the extension

    """
    if path is None:
      path = os.path.join(os.getcwd(), "outs", "w3schools")
    else:
      path = os.path.join(path, "w3schools")
    if not os.path.exists(path):
      os.makedirs(path, exist_ok=True)
    if name is None:
      name = int(time.time())
    file_path = os.path.join(path, "%s.html" % name)
    with open(file_path, "w") as f:
      f.write(self._repr_html_())
    return file_path

  def codepen(self, path=None, name=None):
    """
    Description:
    ------------

    Usage:
    ------

    Related Pages:
    --------------
    https://codepen.io/

    Attributes:
    ----------
    :param path: The path in which the output files will be created
    :param name: The filename without the extension

    TODO Try to add the prefill
    https://blog.codepen.io/documentation/api/prefill/

    :return: The file path
    """
    self.jsfiddle(path, name, framework="codepen")

  def jsfiddle(self, path=None, name=None, framework="jsfiddle"):
    """
    Description:
    ------------
    Produce files which can be copied directly to https://jsfiddle.net in order to test the results and perform changes.

    The output is always in a sub directory jsfiddle

    Usage:
    ------

    Related Pages:
    --------------
    https://jsfiddle.net/

    Attributes:
    ----------
    :param path: The path in which the output files will be created
    :param name: The filename without the extension

    :return: The file path
    """
    if path is None:
      path = os.path.join(os.getcwd(), "outs", framework)
    else:
      path = os.path.join(path, framework)
    if not os.path.exists(path):
      os.makedirs(path, exist_ok=True)
    if os.path.exists(path):
      if name is None:
        name = int(time.time())
      results = self._to_html_obj(content_only=True)
      # For the JavaScript builders
      with open(os.path.join(path, "%s.js" % name), "w") as f:
        f.write(results["jsFrgs"])

      # FOr all the doms and imports
      with open(os.path.join(path, "%s.html" % name), "w") as f:
        f.write("%s\n" % results["cssImports"])
        f.write("%s\n" % results["jsImports"])
        f.write(results["content"])

      # For the CSS styles
      with open(os.path.join(path, "%s.css" % name), "w") as f:
        f.write(results["cssStyle"])
    return path

  def html_file(self, path=None, name=None):
    """
    Description:
    ------------
    Function used to generate a static HTML page for the report

    Usage:
    ------

    Attributes:
    ----------
    :param path: The path in which the output files will be created
    :param name: The filename without the extension

    :return: The file full path
    """
    if path is None:
      path = os.path.join(os.getcwd(), "outs")
    if not os.path.exists(path):
      os.makedirs(path, exist_ok=True)
    if name is None:
      name = int(time.time())
    file_path = os.path.join(path, "%s.html" % name)
    with open(file_path, "w") as f:
      htmlParts, cssParts = [], {}
      for objId in self._report.content:
        if self._report.htmlItems[objId].inReport:
          htmlParts.append(self._report.htmlItems[objId].html())
        cssParts.update(self._report.htmlItems[objId].style.get_classes_css())
      body = str(self._report.body.set_content(self._report, "\n".join(htmlParts)))
      results = self._to_html_obj(htmlParts, cssParts)
      results['body'] = body
      results['header'] = self._report.headers
      f.write(HtmlTmplBase.STATIC_PAGE % results)
    return file_path

  def markdown_file(self, path=None, name=None):
    """
    Description:
    ------------
    Writes a Markdown file from the report object

    Attributes:
    ----------
    :param path: The path in which the output files will be created
    :param name: The filename without the extension

    :return: The file path
    """
    if path is None:
      path = os.path.join(os.getcwd(), "outs", "markdowns")
    else:
      path = os.path.join(path, "markdowns")
    if not os.path.exists(path):
      os.makedirs(path, exist_ok=True)
    if os.path.exists(path):
      if name is None:
        name = "md_%s.amd" % int(time.time())

      file_Path = os.path.join(path, name)
      with open(file_Path, "w") as f:
        for objId in self._report.content:
          html_obj = self._report.htmlItems[objId]
          if hasattr(html_obj, "to_markdown"):
            f.write("%s\n" % html_obj.to_markdown(html_obj.vals))
      return file_Path

  def str(self):
    """

    :return:
    """

  def pdf(self):
    """

    :return:
    """
    self.word()
    # Then save the file to pdf

  def word(self, path=None, name=None):
    """
    Description:
    ------------
    Writes the result to an Word document

    Attributes:
    ----------
    :param path: The path in which the output files will be created
    :param name: The filename without the extension

    :return:
    """
    from docx import Document
    from docx.shared import RGBColor

    if path is None:
      path = os.path.join(os.getcwd(), "outs", "office")
    else:
      path = os.path.join(path, "office")
    if not os.path.exists(path):
      os.makedirs(path, exist_ok=True)
    timestamp = time.strftime("%Y%m%d_%H%M%S", time.gmtime())
    name = name or 'word_%s.docx' % timestamp
    document = Document()
    for objId in self._report.content:
      if self._report.htmlItems[objId].inReport:
        try:
          self._report.htmlItems[objId].to_word(document)
        except Exception as err:
          error_title = document.add_heading().add_run("Error")
          error_title.font.color.rgb = RGBColor(255, 0, 0)
          error_title.font.italic = True
          error_paragraph = document.add_paragraph().add_run((str(err)))
          error_paragraph.font.color.rgb = RGBColor(255, 0, 0)
          error_paragraph.font.italic = True
    document.save(os.path.join(path, name))
    return name

  def excel(self, path=None, name=None):
    """
    Description:
    ------------
    Writes the result to an Excel document

    Attributes:
    ----------
    :param path: The path in which the output files will be created
    :param name: The filename without the extension

    :return:
    """
    xls = requires("xlsxwriter", reason='Missing Package', install='xlsxwriter', source_script=__file__)

    if path is None:
      path = os.path.join(os.getcwd(), "outs", "office")
    else:
      path = os.path.join(path, "office")
    if not os.path.exists(path):
      os.makedirs(path, exist_ok=True)
    timestamp = time.strftime("%Y%m%d_%H%M%S", time.gmtime())
    name = name or '%s_%s.xlsx' % (self._report.run.script_name, timestamp)
    workbook = xls.Workbook(os.path.join(path, name))
    worksheet = workbook.add_worksheet()
    cursor = {'row': 0, 'col': 0}
    for objId in self._report.content:
      if self._report.htmlItems[objId].inReport:
        try:
          self._report.htmlItems[objId].to_xls(workbook, worksheet, cursor)
        except Exception as err:
          cell_format = workbook.add_format({'bold': True, 'font_color': 'red'})
          worksheet.write(cursor['row'], 0, str(err), cell_format)
          cursor['row'] += 2
    workbook.close()
    return name

  def power_point(self):
    """
    Description:
    ------------

    :return:
    """

  @property
  def browser(self):
    """
    Description:
    ------------
    This module will require the package webbrowser.
    It will allow outputs to be created directly in the webpages (without using intermediary text files
    """
    return OutBrowsers(self)
