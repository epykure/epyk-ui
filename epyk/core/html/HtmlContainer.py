
from epyk.core.html import Html
from epyk.core.html import Defaults
from epyk.core.html.options import OptPanel
from epyk.core.html.options import OptText
#
from epyk.core.js import JsUtils
from epyk.core.js.html import JsHtml
from epyk.core.js.html import JsHtmlPanels

# The list of CSS classes#
from epyk.core.css import Defaults as css_defaults
from epyk.core.css.styles import GrpClsContainer


class Panel(Html.Html):
  name, category, callFnc = 'Panel', 'Layouts', 'panel'

  def __init__(self, report, htmlObj, title, color, width, height, htmlCode, helper, options, profile):
    if isinstance(htmlObj, list) and htmlObj:
      for obj in htmlObj:
        if hasattr(obj, 'inReport'):
          obj.inReport = False
    elif htmlObj is not None and hasattr(htmlObj, 'inReport'):
      htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    component = []
    if title is not None:
      self.title = report.ui.title(title)
      self.title.inReport = False
      component.append(self.title)
    container = report.ui.div(htmlObj)
    container.inReport = False
    component.append(container)
    self.add_helper(helper)
    super(Panel, self).__init__(report, component, code=htmlCode, profile=profile,
                                css_attrs={"color": color, "width": width, "height": height})
    container.set_attrs(name="name", value="panel_%s" % self.htmlId)

  @property
  def style(self):
    """
    Description:
    ------------

    :rtype: GrpClsContainer.ClassDiv
    """
    if self._styleObj is None:
      self._styleObj = GrpClsContainer.ClassDiv(self)
    return self._styleObj

  def __add__(self, htmlObj):
    """ Add items to a container """
    htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    self.val.append(htmlObj)
    return self

  @property
  def dom(self):
    """
    Description:
    ------------
    Javascript Functions

    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object

    :rtype: JsHtmlPanels.JsHtmlPanel
    """
    if self._dom is None:
      self._dom = JsHtmlPanels.JsHtmlPanel(self, report=self._report)
    return self._dom

  def __str__(self):
    str_div = "".join([v.html() if hasattr(v, 'html') else str(v) for v in self.val])
    return "<div %s>%s</div>%s" % (self.get_attrs(pyClassNames=self.style.get_classes()), str_div, self.helper)


class PanelSplit(Html.Html):
  __reqJs, __reqCss = ['jqueryui'], ['jqueryui']
  name, category, callFnc = 'Panel Split', 'Layouts', 'panelsplit'

  def __init__(self, report, width, height, left_width, left_obj, right_obj, resizable, helper, profile):
    super(PanelSplit, self).__init__(report, None, css_attrs={"width": width, "height": height, 'white-space': 'nowrap'}, profile=profile)
    self.left_width, self.resizable = left_width, resizable
    if left_obj is not None:
      self.left(left_obj)
    if right_obj is not None:
      self.right(right_obj)
    self.css_left = {'flex': '0 0 auto', 'overflow': 'auto', 'padding': '5px', 'min-width': '100px', 'width': "%s%s" % (self.left_width[0], self.left_width[1]),
                     'white-space': 'nowrap'}
    self.css_right = {'flex': '0 1 auto', 'overflow': 'auto', 'padding': '5px', 'width': '100%', 'background': self._report.theme.greys[0],
                      'border-left': '3px solid %s' % self._report.theme.success[1]}
    self.css({'display': 'flex', 'flex-direction': 'row', 'overflow': 'hidden', 'xtouch-action': 'none'})

  def left(self, html_obj):
    """
    Description:
    ------------
    Add the left component to the panel

    Attributes:
    ----------
    :param html_obj:
    """
    html_obj.inReport = False
    self.html_left = html_obj
    return self

  def right(self, html_obj):
    """
    Description:
    ------------
    Add the right component to the panel

    Attributes:
    ----------
    :param html_obj:
    """
    html_obj.inReport = False
    self.html_right = html_obj
    return self

  def __str__(self):
    self._report._props.setdefault('js', {}).setdefault("builders", []).extend([
      '$("#%(htmlId)s_left").resizable({handleSelector: ".splitter", resizeHeight: false});' % {'htmlId': self.htmlId},
      '$("#%(htmlId)s_right").resizable({handleSelector: ".splitter-horizontal", resizeWidth: true})' % {
        'htmlId': self.htmlId}])
    return '''
      <div %(attrs)s>
        <div style="%(css_left)s" id="%(htmlId)s_left" class="panel-left">%(left)s</div>
        <div style="%(css_right)s" id="%(htmlId)s_right" class="panel-right">%(right)s</div>
      </div>
      ''' % {"attrs": self.get_attrs(pyClassNames=self.style.get_classes()), "htmlId": self.htmlId, 'left': self.html_left.html(),
             'right': self.html_right.html(), 'css_left': css_defaults.inline(self.css_left), 'css_right': css_defaults.inline(self.css_right)}


class PanelSlide(Panel):
  __reqCss, __reqJs = ['font-awesome'], ['font-awesome']
  name, category, callFnc = 'Slide Panel', 'Panels', 'slide'

  def __init__(self, report, htmlObj, title, color, width, height, htmlCode, helper, options, profile):
    super(PanelSlide, self).__init__(report, htmlObj, None, color, width, height, htmlCode, helper, options, profile)
    self.add_helper(helper)
    self.icon = self._report.ui.icon("").css({"display": 'inline-block', 'margin': '0 2px',
                                              'line-height': "%spx" % Defaults.LINE_HEIGHT, 'font-size': "%spx" % Defaults.BIG_ICONS})
    self.text = self._report.ui.title(title).css({"display": 'inline-block', 'margin': 0})
    self.title = self._report.ui.div([self.icon, self.text])
    self.title.inReport = False
    self.title.css({"cursor": 'pointer', "padding": "0 2px 0 0"})
    self._vals, self.__clicks = [self.title] + self._vals, []
    self.__options = OptPanel.OptionPanelSliding(report, options)

  @property
  def options(self):
    """
    Description:
    ------------

    :rtype: OptPanel.OptionPanelSliding
    """
    return self.__options

  def click(self, jsFncs, profile=False):
    """
    Description:
    ------------
    Event added to the title bar.
    This will be triggered first

    Attributes:
    ----------
    :param jsFncs:
    :param profile:
    """
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    self.__clicks = jsFncs
    return self

  def __add__(self, htmlObj):
    """ Add items to a container """
    self.val[1] += htmlObj
    return self

  def __str__(self):
    if self.options.expanded:
      icon_change = self.options.icon_closed
      icon_current = self.options.icon_expanded
      self.icon.set_icon(self.options.icon_expanded)
    else:
      icon_change = self.options.icon_expanded
      icon_current = self.options.icon_closed
      self._vals[1].style.css.display = 'none'
      self.icon.set_icon(self.options.icon_closed)
    self.title.click(self.__clicks + [
      self._report.js.getElementsByName("panel_%s" % self.htmlId).first.toggle(),
      self.icon.dom.switchClass(icon_current, icon_change)])
    str_div = "".join([v.html() if hasattr(v, 'html') else str(v) for v in self.val])
    return "<div %s>%s</div>%s" % (self.get_attrs(pyClassNames=self.style.get_classes()), str_div, self.helper)


class Div(Html.Html):
  name, category, callFnc = 'Simple Container', 'Layouts', 'div'

  def __init__(self, report, htmlObj, label, color, width, icon, height, editable, align, padding, htmlCode, tag,
               helper, profile):
    if isinstance(htmlObj, list) and htmlObj:
      newHtmlObj = []
      for obj in htmlObj:
        if isinstance(obj, list) and obj:
          newHtmlObj.append(report.ui.div(obj, label, color, width, icon, height, editable, align, padding,
                                          htmlCode, tag, helper, profile))
        else:
          newHtmlObj.append(obj)
        if hasattr(newHtmlObj[-1], 'inReport'):
          newHtmlObj[-1].inReport = False
          #if newHtmlObj[-1].css("display") != 'None':
          #  newHtmlObj[-1].css({"display": 'inline-block'})
      htmlObj = newHtmlObj
    elif htmlObj is not None and hasattr(htmlObj, 'inReport'):
      htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    super(Div, self).__init__(report, htmlObj, code=htmlCode, css_attrs={"color": color, "width": width, "height": height},
                              profile=profile)
    self.tag = tag
    # Add the component predefined elements
    self.add_icon(icon)
    self.add_label(label)
    self.add_helper(helper)

    self.css({'text-align': align})
    if padding is not None:
      self.css('padding', '%s' % padding)
    if editable:
      self.set_attrs(name='contenteditable', value="true")
      self.css('overflow', 'auto')

  @property
  def dom(self):
    """
    Description:
    ------------
    Javascript Functions

    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object

    :rtype: JsHtml.JsHtmlRich
    """
    if self._dom is None:
      self._dom = JsHtml.JsHtmlRich(self, report=self._report)
    return self._dom

  def __add__(self, htmlObj):
    """ Add items to a container """
    htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    if not isinstance(self.val, list):
      self.val = [self.val]
    self.val.append(htmlObj)
    return self

  def __getitem__(self, i):
    return self.val[i]

  @property
  def style(self):
    """

    :rtype: GrpClsContainer.ClassDiv
    """
    if self._styleObj is None:
      self._styleObj = GrpClsContainer.ClassDiv(self)
    return self._styleObj

  def build(self, data=None, options=None, profile=False):
    if isinstance(data, dict):
      # check if there is no nested HTML components in the data
      js_data = "{%s}" % ",".join(["%s: %s" % (k, JsUtils.jsConvertData(v, None)) for k, v in data.items()])
    else:
      js_data = JsUtils.jsConvertData(data, None)
    options, js_options = options or {}, []
    for k, v in options.items():
      if isinstance(v, dict):
        row = ["%s: %s" % (s_k, JsUtils.jsConvertData(s_v, None)) for s_k, s_v in v.items()]
        js_options.append("%s: {%s}" % (k, ", ".join(row)))
      else:
        js_options.append("%s: %s" % (k, JsUtils.jsConvertData(v, None)))
    return "%s.innerHTML = %s" % (self.dom.varId, js_data) #, "{%s}" % ",".join(js_options))

  def __str__(self):
    str_div = "".join([v.html() if hasattr(v, 'html') else str(v) for v in self.val])
    return "<div %s>%s</div>%s" % (self.get_attrs(pyClassNames=self.style.get_classes()), str_div, self.helper)

  # -----------------------------------------------------------------------------------------
  #                                    EXPORT OPTIONS
  # -----------------------------------------------------------------------------------------
  def to_word(self, document):
    if isinstance(self.vals, list):
      for val in self.vals:
        if hasattr(val, 'inReport'):
          val.to_word(document)
    else:
      if hasattr(self.vals, 'inReport'):
        self.vals.to_word(document)


class Td(Html.Html):
  name, category, callFnc = 'Column', 'Layouts', 'col'

  def __init__(self, report, htmlObjs, header, position, width, height, align, profile):
    self.position, self.rows_css, self.row_css_dflt, self.header = position, {}, {}, header
    super(Td, self).__init__(report, [], css_attrs={"width": width, "height": height, 'white-space': 'nowrap'}, profile=profile)
    if htmlObjs is not None:
      for htmlObj in htmlObjs:
        self.__add__(htmlObj)

  def __add__(self, htmlObj):
    """ Add items to a container """
    if hasattr(htmlObj, 'inReport'):
      htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    self.val.append(htmlObj)
    return self

  def colspan(self, i):
    """
    Description:
    ------------
    The colspan attribute defines the number of columns a cell should span.

    Related Pages:
    --------------
    https://www.w3schools.com/tags/att_td_colspan.asp

    Attributes:
    ----------
    :param i:
    """
    self.attr['colspan'] = i
    return self

  def rowspan(self, i):
    """
    Description:
    ------------
    The rowspan attribute specifies the number of rows a cell should span.

    Related Pages:
    --------------
    https://www.w3schools.com/tags/att_td_rowspan.asp

    Attributes:
    ----------
    :param i:
    """
    self.attr['rowspan'] = i
    return self

  def __getitem__(self, i):
    return self.val[i]

  def __str__(self):
    content = [htmlObj.html() if hasattr(htmlObj, 'inReport') else str(htmlObj) for htmlObj in self.val]
    if self.header:
      return '<th %s>%s</th>' % (self.get_attrs(pyClassNames=self.style.get_classes()), "".join(content))

    return '<td %s>%s</td>' % (self.get_attrs(pyClassNames=self.style.get_classes()), "".join(content))


class Tr(Html.Html):
  name, category, callFnc = 'Column', 'Layouts', 'col'

  def __init__(self, report, htmlObjs, header, position, width, height, align, options, profile):
    self.position, self.header = position, header
    super(Tr, self).__init__(report, [], css_attrs={"width": width, "height": height, 'text-align': align}, profile=profile)
    if htmlObjs is not None:
      for htmlObj in htmlObjs:
        self.__add__(htmlObj)
    self.style.justify_content = self.position

  def __add__(self, htmlObj):
    """ Add items to a container """
    if not isinstance(htmlObj, Td):
      if not isinstance(htmlObj, list):
        htmlObj = [htmlObj]
      htmlObj = Td(self._report, htmlObj, self.header, None, (None, "%"), (None, "%"), None, False)
    self.val.append(htmlObj)
    return self

  @property
  def dom(self):
    """
    Javascript Functions

    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object

    :rtype: JsHtmlPanels.JsHtmlTr
    """
    if self._dom is None:
      self._dom = JsHtmlPanels.JsHtmlTr(self, report=self._report)
    return self._dom

  def __getitem__(self, i):
    """
    Description:
    ------------
    Return the internal column in the row for the given index

    Attributes:
    ----------
    :param i: the column index

    :rtype: Td
    """
    return self.val[i]

  def __str__(self):
    cols = [htmlObj.html() for i, htmlObj in enumerate(self.val)]
    return '<tr %s>%s</tr>' % (self.get_attrs(pyClassNames=self.style.get_classes()), "".join(cols))


class Caption(Html.Html):

  def __init__(self, report, text, color, align, width, height, htmlCode, tooltip, options, profile):
    super(Caption, self).__init__(report, text, css_attrs={"width": width, "height": height, "color": color, 'text-align': align},
                               code=htmlCode, profile=profile)
    self.__options = OptText.OptionsText(self, options)
    if tooltip is not None:
      self.tooltip(tooltip)

  @property
  def options(self):
    """
    Description:
    ------------
    Property to set all the possible object for a button

    :rtype: OptText.OptionsText
    """
    return self.__options

  def __str__(self):
    val = self._report.py.markdown.all(self.val) if self.options.showdown else self.val
    return '<caption %s>%s</caption>' % (self.get_attrs(pyClassNames=self.style.get_classes()), self.val)


class TSection(Html.Html):

  def __init__(self, report, type, rows=None):
    super(TSection, self).__init__(report, [])
    self.__section = type
    if rows is not None:
      for row in rows:
        self.__add__(row)

  def __getitem__(self, i):
    """
    Description:
    ------------

    Attributes:
    ----------
    :param i: Integer. The internal row based on the index

    :rtype: Tr
    """
    return self.val[i]

  def __add__(self, row_data):
    """ Add items to a container """
    if not isinstance(row_data, Tr):
      row_data = Tr(self._report, row_data, self.__section == 'thead', None, (100, "%"), (100, "%"), 'center', None, False)

    self.val.append(row_data)
    return self

  def __str__(self):
    cols = "".join([htmlObj.html() for i, htmlObj in enumerate(self.val)])
    return '<%(section)s %(attr)s>%(cols)s</%(section)s>' % {'section': self.__section, 'cols': cols,
                  'attr': self.get_attrs(pyClassNames=self.style.get_classes())}


class Table(Html.Html):
  name, category, callFnc = 'Table', 'Layouts', 'table'

  def __init__(self, report, rows, width, height, align, helper, options, profile):
    super(Table, self).__init__(report, [], css_attrs={"width": width, "height": height, 'table-layout': 'auto',
            'white-space': 'nowrap', 'border-collapse': 'collapse', 'box-sizing': 'border-box'}, profile=profile)
    self.add_helper(helper, css={"float": "none", "margin-left": "5px"})
    self.header, self.body, self.footer = TSection(self._report, 'thead'), TSection(self._report, 'tbody'), TSection(self._report, 'tfoot')
    self.caption = None
    if rows is not None:
      for row in rows:
        self.__add__(row)

  def __add__(self, row_data):
    """ Add items to a container """
    if isinstance(row_data, Tr):
      row = row_data
    else:
      if not self.header.val and not self.body.val:
        row = Tr(self._report, row_data, True, None, (100, "%"), (100, "%"), 'center', None, False)
      else:
        row = Tr(self._report, row_data, False, None, (100, "%"), (100, "%"), 'left', None, False)
    if row.header:
      self.header += row
    else:
      self.body += row
    return self

  def add_caption(self, text, color=None, align=None, width=(100, "%"), height=(100, "%"), htmlCode=None, tooltip=None,
                  options=None, profile=False):
    """
    Description:
    ------------
    The <caption> tag defines a table caption.

    Related Pages:
    --------------
    https://www.w3schools.com/tags/tag_caption.asp

    Attributes:
    ----------
    :param text:
    :param color:
    :param align:
    :param width:
    :param height:
    :param htmlCode:
    :param tooltip:
    :param options:
    :param profile:
    """
    self.caption = Caption(self._report, text, color, align, width, height, htmlCode, tooltip, options, profile)
    return self.caption

  def get_header(self, i=0):
    """
    Description:
    ------------

    Attributes:
    ----------
    :param i: Integer.

    :return:
    """
    return self.header.val[i]

  def get_footer(self, i=0):
    """
    Description:
    ------------

    Attributes:
    ----------
    :param i: Integer.

    :return:
    """
    return self.footer.val[i]

  def __getitem__(self, i):
    """
    Description:
    ------------

    Attributes:
    ----------
    :param i: Integer. The internal row based on the index

    :rtype: Tr
    """
    if not self.body.val:
      return []

    return self.body.val[i]

  def __str__(self):
    caption = "" if self.caption is None else self.caption.html()
    return '<table %s>%s%s%s%s</table>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), caption, self.header.html(), self.body.html(), self.footer.html(), self.helper)


class Col(Html.Html):
  name, category, callFnc = 'Column', 'Layouts', 'col'

  def __init__(self, report, htmlObjs, position, width, height, align, helper, profile):
    self.position,  self.rows_css, self.row_css_dflt = position, {}, {}
    super(Col, self).__init__(report, [], css_attrs={"width": width, "height": height}, profile=profile)
    if htmlObjs is not None:
      for htmlObj in htmlObjs:
        self.__add__(htmlObj)
    if align == "center":
      self.css({'margin': 'auto', 'display': 'inline-block', 'text-align': 'center'})
    else:
      self.css({'display': 'inline-block'})
    self.attr["class"].add('col')
    self.style.justify_content = self.position

  def __add__(self, htmlObj):
    """
    Description:
    ------------
    Add items to a container

    Attributes:
    ----------
    :param htmlObj:

    :return:
    """
    if not hasattr(htmlObj, 'inReport'):
      htmlObj = self._report.ui.div(htmlObj)
    htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    self.val.append(htmlObj)
    return self

  def build(self, data=None, options=None, profile=False):
    """
    Description:
    ------------

    Attributes:
    ----------
    :param data:
    :param options:
    :param profile:
    """
    return self.val[0].build(data, options, profile)

  def set_size(self, n):
    """
    Description:
    ------------
    Set the column size

    Usage:
    ------
    ps = rptObj.ui.layouts.grid()
    ps += [rptObj.ui.text("test %s" % i) for i in range(5)]
    ps[0][0].set_size(10)

    Attributes:
    ----------

    """
    self.attr["class"].add("col-%s" % n)
    return self

  def __getitem__(self, i):
    return self.val[i]

  def __str__(self):
    content = [htmlObj.html() for htmlObj in self.val]
    return '<div %s>%s</div>' % (self.get_attrs(pyClassNames=self.style.get_classes()), "".join(content))

  # -----------------------------------------------------------------------------------------
  #                                    EXPORT OPTIONS
  # -----------------------------------------------------------------------------------------
  def to_word(self, document):
    for i, htmlObj in enumerate(self.vals):
      htmlObj.to_word(document)

  def to_xls(self, workbook, worksheet, cursor):
    for i, htmlObj in enumerate(self.vals):
      try:
        htmlObj.to_xls(workbook, worksheet, cursor)
      except Exception as err:
        cell_format = workbook.add_format({'bold': True, 'font_color': 'red'})
        worksheet.write(cursor['row'], 0, str(err), cell_format)
        cursor['row'] += 2


class Row(Html.Html):
  name, category, callFnc = 'Column', 'Layouts', 'col'
  __reqCss, __reqJs = ['bootstrap'], ['bootstrap']

  def __init__(self, report, htmlObjs, position, width, height, align, helper, profile):
    self.position = position
    super(Row, self).__init__(report, [], css_attrs={"width": width, "height": height}, profile=profile)
    if htmlObjs is not None:
      for htmlObj in htmlObjs:
        self.__add__(htmlObj)
    self.attr["class"].add('row')
    self.attr["class"].add('no-gutters')
    self.style.justify_content = self.position

  @property
  def dom(self):
    """
    Description:
    ------------
    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object

    :rtype: JsHtmlPanels.JsHtmlRow
    """
    if self._dom is None:
      self._dom = JsHtmlPanels.JsHtmlRow(self, report=self._report)
    return self._dom

  def __add__(self, htmlObj):
    """ Add items to a container """
    if not isinstance(htmlObj, Col):
      if not isinstance(htmlObj, list):
        htmlObj = [htmlObj]
      htmlObj = self._report.ui.layouts.col(htmlObj)
    htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    self.val.append(htmlObj)
    return self

  def __getitem__(self, i):
    """
    Return the internal column in the row for the given index

    :param i: the column index
    :rtype: Col
    """
    return self.val[i]

  def __str__(self):
    cols = []
    for i, htmlObj in enumerate(self.val):
      cols.append(htmlObj.html())
    return '<div %s>%s</div>' % (self.get_attrs(pyClassNames=self.style.get_classes()), "".join(cols))


class Grid(Html.Html):
  name, category, callFnc = 'Grid', 'Layouts', 'grid'
  __reqCss, __reqJs = ['bootstrap'], ['bootstrap']

  def __init__(self, report, rows, width, height, align, helper, profile):
    super(Grid, self).__init__(report, [], css_attrs={"width": width, "height": height}, profile=profile)
    self.css({'overflow-x': 'hidden', 'padding': 0})
    self.attr["class"].add("container-fluid")
    if align == 'center':
      self.css({'margin': 'auto'})
    if rows is not None:
      for row in rows:
        self.__add__(row)

  def __add__(self, row_data):
    """ Add items to a container """
    if isinstance(row_data, Row):
      row = row_data
    else:
      row = self._report.ui.layouts.row()
      for htmlObjWithDim in row_data:
        if isinstance(htmlObjWithDim, tuple):
          htmlObj, dim = htmlObjWithDim
        else:
          htmlObj, dim = htmlObjWithDim, None
        row += htmlObj
        if dim is not None:
          row[-1].attr["class"].add("col-%s" % dim)
    row.inReport = False
    self.val.append(row)
    #self.colsAlign.append("left")
    return self

  def __getitem__(self, i):
    """
    Description:
    ------------

    Attributes:
    ----------
    :param i: Integer. The internal row based on the index

    :rtype: Row
    """
    return self.val[i]

  @property
  def dom(self):
    """
    Description:
    ------------
    Javascript Functions

    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object

    :rtype: JsHtmlPanels.JsHtmlGrid
    """
    if self._dom is None:
      self._dom = JsHtmlPanels.JsHtmlGrid(self, report=self._report)
    return self._dom

  def resize(self):
    """
    Description:
    ------------
    For the resizing of the space for the containers.

    This will rescale based on the number of items and the fact that the max per row is 12
    It will update the colsDim list
    """
    max_size = int(12 / len(self.colsDim))
    self.colsDim = [max_size] * len(self.colsDim)
    return self

  def __str__(self):
    rows = [htmlObj.html() for htmlObj in self.val]
    return '<div %s>%s</div>' % (self.get_attrs(pyClassNames=self.style.get_classes()), "".join(rows))


  # -----------------------------------------------------------------------------------------
  #                                    EXPORT OPTIONS
  # -----------------------------------------------------------------------------------------
  def to_word(self, document):
    for i, htmlObj in enumerate(self.vals):
      try:
        htmlObj.to_word(document)
      except Exception as err:
        from docx.shared import RGBColor

        errotTitle = document.add_heading().add_run("Error")
        errotTitle.font.color.rgb = RGBColor(255, 0, 0)
        errotTitle.font.italic = True
        errorParagraph = document.add_paragraph().add_run((str(err)))
        errorParagraph.font.color.rgb = RGBColor(255, 0, 0)
        errorParagraph.font.italic = True

  def to_xls(self, workbook, worksheet, cursor):
    for i, htmlObj in enumerate(self.vals):
      try:
        htmlObj.to_xls(workbook, worksheet, cursor)
      except Exception as err:
        cell_format = workbook.add_format({'bold': True, 'font_color': 'red'})
        worksheet.write(cursor['row'], 0, str(err), cell_format)
        cursor['row'] += 2

  # -----------------------------------------------------------------------------------------
  #                                    MARKDOWN SECTION
  # -----------------------------------------------------------------------------------------
  @staticmethod
  def matchMarkDown(val): return True if val == "[" else None

  @classmethod
  def convertMarkDown(cls, val, regExpResult, report=None):
    pass

  @classmethod
  def jsMarkDown(self, vals):
    return '''
      '''


class Tabs(Html.Html):
  name, category, callFnc = 'Tabs', 'Layouts', 'tabs'

  def __init__(self, report, color, width, height, htmlCode, helper, options, profile):
    super(Tabs, self).__init__(report, "", code=htmlCode, css_attrs={"width": width, "height": height, 'color': color}, profile=profile)
    self.__panels, self.__panel_objs, self.__selected = [], {}, None
    self.tabs_name, self.panels_name = "button_%s" % self.htmlId, "panel_%s" % self.htmlId
    self.tabs_container = self._report.ui.div([])
    self.tabs_container.inReport = False
    self.add_helper(helper)
    self.__options = OptPanel.OptionPanelTabs(report, options)

  @property
  def options(self):
    """
    Description:
    ------------

    :rtype: OptPanel.OptionPanelTabs
    """
    return self.__options

  @property
  def dom(self):
    """
    Description:
    ------------
    Javascript Functions

    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object

    :rtype: JsHtmlPanels.JsHtmlTabs
    """
    if self._dom is None:
      self._dom = JsHtmlPanels.JsHtmlTabs(self, report=self._report)
    return self._dom

  def __getitem__(self, name):
    return self.__panel_objs[name]

  def select(self, name):
    """
    Description:
    ------------

    Attributes:
    ----------
    :param name:
    """
    self.__selected = name
    return self

  def panel(self, name):
    """
    Description:
    ------------
    Get the panel obkect

    Attributes:
    ----------
    :param name: String. The tab name

    :rtype: Div
    """
    return self.__panel_objs[name]["content"]

  def tab(self, name):
    """
    Description:
    ------------
    Get the tab container

    Attributes:
    ----------
    :param name: String. The tab name

    :rtype: Div
    """
    return self.__panel_objs[name]["tab"][0]

  def add_panel(self, name, div, icon=None, selected=False, css_tab=None, css_tab_clicked=None):
    """
    Description:
    ------------

    Attributes:
    ----------
    :param name:
    :param div:
    :param selected: Boolean. Flag to set the selected panel
    :param css_tab:
    """
    div.css({"display": 'none'})
    div.inReport = False
    div.set_attrs(name="name", value=self.panels_name)
    self.__panels.append(name)
    if icon is not None:
      tab = self._report.ui.div([
        self._report.ui.icon(icon).css({"display": 'block', "width": '100%', "font-size": css_defaults.font(4)}),
        name], width=("100", "px"))
    else:
      tab = self._report.ui.div(name, width=("100", "px"))
    tab_style = self.options.tab_style(name, css_tab)
    tab_style_clicked = self.options.tab_clicked_style(name, css_tab_clicked)
    tab.css(tab_style).css({"padding": '5px 0'})
    tab.set_attrs(name="name", value=self.tabs_name)
    tab.set_attrs(name="data-index", value=len(self.__panels) - 1)
    tab_container = self._report.ui.div(tab, width=("100", "px"))
    tab_container.inReport = False
    tab_container.css({'display': 'inline-block'})
    css_cls_name = None
    tab.click([
      self.dom.deselect_tabs(),
      tab.dom.setAttribute("data-selected", True).r,
      self._report.js.getElementsByName(self.panels_name).all([
        #self._report.js.getElementsByName(self.tabs_name).all([
        #  self._report.js.data.all.element.css(self.options.tab_not_clicked_style(name))]),
        tab.dom.css(tab_style_clicked),
        self._report.js.data.all.element.hide(),
        tab_container.dom.toggleClass(css_cls_name, propagate=True) if css_cls_name is not None else "",
        div.dom.show()])])
    tab.inReport = False
    self.__panel_objs[name] = {"tab": tab_container, "content": div}
    if selected:
      self.__selected = name
    return self

  def __str__(self):
    if self.__selected is not None:
      self.__panel_objs[self.__selected]["content"].style.css.display = 'block'
      self.__panel_objs[self.__selected]["tab"][0].css(self.options.tab_clicked_style(self.__selected))
      self.__panel_objs[self.__selected]["tab"][0].attr["data-selected"] = 'true'
    content = []
    for p in self.__panels:
      self.tabs_container += self.__panel_objs[p]["tab"]
      content.append(self.__panel_objs[p]["content"].html())
    return "<div %s>%s%s</div>%s" % (self.get_attrs(pyClassNames=self.style.get_classes()), self.tabs_container.html(), "".join(content), self.helper)


class IFrame(Html.Html):
  name, category, callFnc = 'IFrame', 'Container', 'iframe'

  def __init__(self, report, url, width, height, helper, profile):
    super(IFrame, self).__init__(report, url, css_attrs={"width": width, "height": height}, profile=profile)
    self.css({"overflow-x": 'hidden'})
    self.add_helper(helper)

  @property
  def _js__builder__(self):
    return 'htmlObj.src = data'

  def __str__(self):
    return "<iframe src='%s' %s frameborder='0' scrolling='no'></iframe>%s" % (self.val, self.get_attrs(pyClassNames=self.style.get_classes()), self.helper)


class Dialog(Html.Html):
  name, category, callFnc = 'DialogMenu', 'Layouts', 'dialogs'
  __reqCss, __reqJs = ['jqueryui'], ['jqueryui']

  def __init__(self, report, recordSet, width, height, helper, profile):
    super(Dialog, self).__init__(report, recordSet, css_attrs={"width": width, 'height': helper}, profile=profile)
    self.css({"border": '2px solid %s' % self._report.theme.greys[3], "display": "block", "position": "absolute",
              "background": self._report.theme.greys[0]})
    # self._report._props.setdefault('js', {}).setdefault("builders", []).append(self.dom.jquery_ui.draggable().toStr())

  @property
  def _js__builder__(self):
    return '''
      var div = jQuery("<div>")
      jQuery(htmlObj).append(div);
      div.dialog({modal: false, title: "rrrr"})'''

  # def jsAdd(self, title='data.event_val', isPyData=False):
  #   if isPyData:
  #     title = json.loads(title)
  #
  #   return '''
  #   var dialogWindow = $("<div>");
  #   var table = $("table");
  #   dialogWindow.append(table);
  #   dialogWindow.append('<input type="text">');
  #   var d = dialogWindow.dialog( { modal: false, title: %(title)s, show: 'puff', fluid: true,
  #       close: function () {$(this).remove()}, appendTo: "#%(htmlId)s", resizable: false,
  #       buttons: [{text: "Close", click: function() { $( this ).dialog("close")} } ]
  #   });
  #   d.parent().draggable({containment: '#%(htmlId)s'});
  #   event.preventDefault()''' % {'title': title, 'htmlId': self.htmlId}

  def __str__(self):
    return "<div %s></div>" % self.get_attrs(pyClassNames=self.style.get_classes())


class IconsMenu(Html.Html):
  name, category, callFnc = 'Icons Menu', 'Layouts', 'menu'
  __reqCss, __reqJs = ['font-awesome'], ['font-awesome']

  def __init__(self, icon_names, report, width, height, htmlCode, helper, profile):
    super(IconsMenu, self).__init__(report, None, width=width, css_attrs={"width": width, "height": height}, code=htmlCode,
                                    profile=profile)
    self._jsActions, self._definedActions = {}, []
    self._icons, self.icon = [], None
    self.css({"margin": "5px 0"})
    for i in icon_names:
      self.add_icon(i)

  def __getitem__(self, i):
    return self._icons[i]

  def add_icon(self, text, css=None, position="after"):
    """
    Add an icon to the HTML object

    Example
    checks.title.add_icon("fas fa-align-center")

    Documentation

    :param text: The icon reference from font awsome website
    :param css: Optional. A dictionary with the CSS style to be added to the component
    :param position:
    :return: The Html object
    """
    if text is not None:
      self._icons.append(self._report.ui.images.icon(text).css({"margin-right": '5px', 'cursor': "pointer"}))
      self.icon = self._icons[-1]
      if position == "before":
        self.prepend_child(self.icon)
      else:
        self.append_child(self.icon)
      #elf.icon.inReport = False
      if css is not None:
        self.icon.css(css)
    return self

  def add_select(self, action, data, width=150):
    options = ["<option>%s</option>" % d for d in data]
    self._jsActions[action] = '<select id="inputState" class="form-control" style="width:%spx;display:inline-block">%s</select>' % (width, "".join(options))
    self._definedActions.append(action)
    return self

  def __str__(self):
    htmlIcons = [htmlDef for action, htmlDef in self._jsActions.items()]
    return "<div %s>%s</div>" % (self.get_attrs(pyClassNames=self.style.get_classes()), "".join(htmlIcons))


class Form(Html.Html):
  name, category, callFnc = 'Generic Form', 'Forms', 'form'

  def __init__(self, report, htmlObjs, helper):
    super(Form, self).__init__(report, [])
    self.style.css.padding = "5px"
    self.add_helper(helper)
    self.__submit, self._has_container = None, False
    for i, htmlObj in enumerate(htmlObjs):
      self.__add__(htmlObj)

  def __getitem__(self, i):
    """
    Return the internal column in the row for the given index

    :param i: the column index
    :rtype: Col
    """
    return self.val[i]

  def __add__(self, htmlObj):
    """ Add items to a container """
    htmlObj.inReport = False # Has to be defined here otherwise it is set too late
    htmlObj.css({'text-align': 'left'})
    self.val.append(htmlObj)
    return self

  def submit(self, method, action="#", text="Submit"):
    if action is not None and method is not None:
      self.attr.update({"action": action, "method": method})

    self.__submit = self._report.ui.button(text).set_attrs({"type": 'submit'})
    self.__submit.inReport = False
    if self._has_container:
      self[0] + self.__submit
    return self

  def __str__(self):
    if self.__submit is None:
      raise Exception("Submit must be defined in a form ")

    str_vals = "".join([i.html() for i in self.val]) if self.val is not None else ""
    return '<form %s>%s%s</form>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), str_vals, self.__submit.html(), self.helper)


class Modal(Html.Html):
  name, category, callFnc = 'Modal Popup',  'Container', 'modal'

  def __init__(self, report, htmlObjs, header, footer, submit, helper):
    """
    Description:
    -----------
    Constructor for the modal item.
    This object is composed of three parts:  a header, which is a row of object, a body which is a column and a footer which is a row
    they all accept collections of html objects and are configurable just like the normal rows and column objects

    """
    super(Modal, self).__init__(report, [])
    self.add_helper(helper)
    self.doSubmit = submit
    if self.doSubmit:
      self.submit = report.ui.buttons.important("Submit").set_attrs({"type": 'submit'})
      self.submit.inReport = False
    self.closeBtn = report.ui.texts.span('&times', width='auto')
    self.closeBtn.css(None, reset=True)
    self.closeBtn.style.add_classes.div.span_close()
    self.closeBtn.click(report.js.getElementById(self.htmlId).css({'display': "none"}))
    self.__header = report.ui.row([])
    self.__header.inReport = False
    if header:
      for obj in header:
        self.__header + obj
    self.__header += self.closeBtn
    if footer:
      for obj in footer:
        self.__footer + obj
    self.__footer = report.ui.row([])
    self.__footer.inReport = False
    self.__body = report.ui.col([]).css({'position': 'relative',  'overflow-y': 'scroll'})
    self.__body.inReport = False
    self.col = report.ui.col([self.__header, self.__body, self.__footer]).css({'width': 'auto'}, reset=True)
    self.col.style.add_classes.div.modal_content()
    self.col.inReport = False
    self.val.append(self.col)
    self.__outOfScopeClose = True
    for htmlObj in htmlObjs:
      self.__add__(htmlObj)

  @property
  def outOfScopeClose(self):
    return self.__outOfScopeClose

  @outOfScopeClose.setter
  def outOfScopeClose(self, val):
    self.__outOfScopeClose = val

  @property
  def style(self):
    """
    Property to the CSS Style of the component

    :rtype: GrpClsButton.ClassButton
    """
    if self._styleObj is None:
      self._styleObj = GrpClsContainer.ClassModal(self)
    return self._styleObj

  @property
  def header(self):
    return self.__header

  @property
  def footer(self):
    return self.__footer

  @property
  def body(self):
    return self.__body

  def show(self):
    return self._report.js.getElementById(self.htmlId).css({'display': 'block'})

  def close(self):
    return self._report.js.getElementById(self.htmlId).css({'display': 'none'})

  def close_on_background(self):
    """
    Description:
    ------------
    Will allow an event to close the modal if a click event is detected anywhere outside the modal
    """
    modal = self._report.js.getElementById(self.htmlId)
    self._report.js.addOnReady(self._report.js.window.events.
                                      addClickListener(self._report.js.if_('event.target == %s' % modal, modal.css({'display': 'none'})),
                                                       subEvents=['event']))

  def __add__(self, htmlObj):
    """ Add items to a container """
    htmlObj.inReport = False # Has to be defined here otherwise it is set too late
    self.__body += htmlObj
    return self

  def __str__(self):
    if self.__outOfScopeClose:
      self.close_on_background()
    str_vals = "".join([i.html() for i in self.val]) if self.val is not None else ""
    self.set_attrs({'css': self.style.css.attrs})
    if self.doSubmit:
      self.col += self.submit
    return '<div %s>%s</div>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), str_vals, self.helper)


class Indices(Html.Html):
  name, category, callFnc = 'Index', 'Panels', 'index'
  __reqCss, __reqJs = ['font-awesome'], ['font-awesome']

  def __init__(self, report, count, width, height, htmlCode, options, profile):
    super(Indices, self).__init__(report, count, css_attrs={"width": width, "height": height}, profile=profile)
    self.items = []
    self.__options = OptPanel.OptionsPanelPoints(report, options)
    for i in range(count):
      div = self._report.ui.div(i, width=(15, "px"))
      div.attr["name"] = self.htmlId
      div.attr["data-position"] = i + 1
      div.css({"display": 'inline-block', "padding": "2px", "text-align": "center"})
      div.css(self.options.div_css)
      div.style.add_classes.div.background_hover()
      div.inReport = False
      self.items.append(div)
    #
    self.first = self._report.ui.icon("fas fa-angle-double-left", width=(20, 'px')).css({"display": 'inline-block'})
    self.first.inReport = False
    self.prev = self._report.ui.icon("fas fa-chevron-left", width=(20, 'px')).css({"display": 'inline-block'})
    self.prev.inReport = False
    self.next = self._report.ui.icon("fas fa-chevron-right", width=(20, 'px')).css({"display": 'inline-block'})
    self.next.inReport = False
    self.last = self._report.ui.icon("fas fa-angle-double-right", width=(20, 'px')).css({"display": 'inline-block'})
    self.last.inReport = False

  @property
  def options(self):
    """

    :rtype: OptPanel.OptionsPanelPoints
    """
    return self.__options

  def __getitem__(self, i):
    return self.items[i]

  def click_item(self, i, jsFncs, profile=False):
    """
    Description:
    ------------

    :param i:
    :param jsFncs:
    :param profile:
    """
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    return self[i].on("click", [
      self[i].dom.by_name.css({"border-bottom": "1px solid %s" % self._report.theme.colors[0]}).r,
      self[i].dom.css({"border-bottom": "1px solid %s" % self.options.background_color})] + jsFncs, profile)

  def __str__(self):
    str_vals = "".join([self.first.html(), self.prev.html()] + [i.html() for i in self.items] + [self.next.html(), self.last.html()])
    return '<div %s>%s</div>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), str_vals, self.helper)


class Points(Html.Html):
  name, category, callFnc = 'Index', 'Panels', 'index'

  def __init__(self, report, count, width, height, htmlCode, options, profile):
    super(Points, self).__init__(report, count, css_attrs={"width": width, "height": height}, profile=profile)
    self.items = []
    self.css({"text-align": "center"})
    self.__options = OptPanel.OptionsPanelPoints(report, options)
    for i in range(count):
      div = self._report.ui.div(self._report.entities.non_breaking_space)
      div.attr["name"] = self.htmlId
      div.attr["data-position"] = i # keep the python indexation
      div.css({"border": "1px solid %s" % self._report.theme.greys[5], "border-radius": "10px", "width": "15px", "height": "15px"})
      div.css(self.options.div_css)
      div.style.add_classes.div.background_hover()
      div.inReport = False
      self.items.append(div)
    self[self.options.selected].css({"background-color": self.options.background_color})

  @property
  def options(self):
    """
    Description:
    ------------

    :rtype: OptPanel.OptionsPanelPoints
    """
    return self.__options

  def on(self, event, jsFncs, profile=False):
    """
    Description:
    ------------
    Add Javascript events to all the items in the component

    Attributes:
    ----------
    :param jsFncs: Array. The Javascript functions
    :param profile:
    :return:
    """
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    str_fnc = JsUtils.jsConvertFncs(jsFncs, toStr=True)
    if event == "click":
      for i in range(len(self.items)):
        self.click_item(i, str_fnc)
    else:
      for i in range(len(self.items)):
        self.on_item(i, event, str_fnc)
    return self

  def on_item(self, i, event, jsFncs, profile=False):
    """
    Description:
    ------------

    Attributes:
    ----------
    :param i: Integer. The item index
    :param event: String. The Javascript event reference
    :param jsFncs: Array. The Javascript functions
    :param profile:
    """
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    return self[i].on(event, [
      'var data = {position: this.getAttribute("data-position")}'] + jsFncs, profile)

  def click_item(self, i, jsFncs, profile=False):
    """
    Description:
    ------------
    Add a click event on a particular item of the component

    Attributes:
    ----------
    :param i: Integer. The item index
    :param jsFncs: Array. The Javascript functions
    :param profile:
    """
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    return self[i].on("click", [
      'var data = {position: this.getAttribute("data-position")}',
      self[i].dom.by_name.css({"background-color": ""}).r,
      self[i].dom.css({"background-color": self.options.background_color})] + jsFncs, profile)

  def __getitem__(self, i):
    return self.items[i]

  def __str__(self):
    str_vals = "".join([i.html() for i in self.items])
    return '<div %s>%s</div>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), str_vals, self.helper)
