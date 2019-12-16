"""
List of all the different templates configurations available for displaying bespoke lists.
This list can be extended and it is easy to test a new configuration by different defining the HTML template in the common list object.
List are standard and very popular HTML objects, please have a look at the below websites if you need further information to manipulate them in your report

"""

import importlib
import os
import inspect
import re
import json

from epyk.core.html import Html

# The list of CSS classes
from epyk.core.css.groups import CssGrpCls
from epyk.core.css.groups import CssGrpClsList


class Li(Html.Html):
  name, category, callFnc = 'Entries', 'Lists', 'list_entries'
  builder_name = False

  def __init__(self, report, text):
    super(Li, self).__init__(report, text)
    self.css({'font-size': 'inherit', 'margin': "1px 5px", 'padding': 0})

  @property
  def no_decoration(self):
    """
    Property to remove the list default style
    """
    self.css({"text-decoration": "none", "list-style-type": 'none'})
    return self

  def add_label(self, text, css=None, position="before", for_=None):
    """
    Add an elementary label component

    Example

    Documentation
    https://www.w3schools.com/tags/tag_label.asp

    :param text: The label content
    :param css: Optional. A dictionary with the CSS style to be added to the component
    :param position:
    :param for_: Specifies which form element a label is bound to
    """
    self.label = ""
    if text is not None:
      dfl_css = {"float": 'none', 'width': 'none'}
      if css is not None:
        dfl_css.update(css)
      self.label = self._report.ui.texts.label(text)
      if for_ is not None:
        # Attach the label to another HTML component based on the ID
        self.label.attr['for'] = for_
      if position == "before":
        self.prepend_child(self.label)
      else:
        self.append_child(self.label)
      self.label.css(dfl_css)
    return self

  def set_html_content(self, htmlObj):
    """
    Set the cell content to be an HTML object

    :param htmlObj: Python HTML object
    :return: self, the cell object to allow the chaining
    """
    htmlObj.inReport = False
    self.innerPyHTML = htmlObj
    return self

  def __str__(self):
    return "<li %s>%s</li>" % (self.get_attrs(pyClassNames=self.defined), self.content)


class List(Html.Html):
  name, category, callFnc = 'List', 'Lists', 'list'
  # The CSS Group attached to this component
  grpCls = CssGrpClsList.CssClassList

  def __init__(self, report, data, size, color, width, height, htmlCode, helper, profile):
    super(List, self).__init__(report, data, width=width[0], widthUnit=width[1], height=height[0],
                               heightUnit=height[1], code=htmlCode, profile=profile)
    self.add_helper(helper)
    self.color = color if color is not None else self.getColor("greys", 9)
    self.css({'font-size': "%s%s" % (size[0], size[1]) if size is not None else 'inherit',
              'padding': 0, 'margin': "1px", 'list-style-position': 'inside'})
    self.items = None
    if len(data) > 0:
      self.set_items()

  def __getitem__(self, i):
    """

    :param i:
    :rtype: Li
    """
    return self.items[i]

  def add_item(self, d):
    """

    :param d:
    :return:
    """
    if self.items is None:
      self.items = []
    li_obj = Li(self._report, d)
    if hasattr(d, 'inReport'):
      d.inReport = False
    self.items.append(li_obj)
    return self

  def set_items(self):
    """
    """
    if self.items is None:
      self.items = []
    for d in self.val:
      li_obj = Li(self._report, d)
      self.items.append(li_obj)
    return self

  @property
  def _js__builder__(self):
    return ''' 
      '''

  def __str__(self):
    self._vals = "".join([i.html() for i in self.items]) if self.items is not None else ""
    #self.builder_name = self.__class__.__name__
    #self._report._props.setdefault('js', {}).setdefault("builders", []).append(self.refresh())
    return "<ul %s>%s</ul>" % (self.get_attrs(pyClassNames=self.defined), self._vals)


class Groups(Html.Html):
  name, category, callFnc = 'Groups', 'Lists', 'groups'
  builder_name = False

  def __init__(self, report, data, categories, size, color, width, height, htmlCode, helper, profile):
    super(Groups, self).__init__(report, [], width=width[0], widthUnit=width[1], height=height[0],
                                 heightUnit=height[1], code=htmlCode, profile=profile)
    self.add_helper(helper)
    self.color = color if color is not None else self.getColor("greys", 9)
    self.css({'font-size': "%s%s" % (size[0], size[1]) if size is not None else 'inherit',
              'margin': "1px", 'padding': '0 2px'})
    self.builder_name, self._lists__map, self._lists__map_index = False, {}, []
    for i, cat in enumerate(categories):
      self.add_list(data[i], cat)

  def __getitem__(self, i):
    return self.val[i]

  def add_list(self, data, category="", size=('inherit', ''), color='inherit', width=(None, 'px'), height=(None, 'px'),
               htmlCode=None, helper=None, profile=False):
    """

    :param data:
    :param category:
    :param size:
    :param color:
    :param width:
    :param height:
    :param htmlCode:
    :param helper:
    :param profile:
    """
    self._lists__map[category] = len(self.val)
    html_li = List(self._report, data, size, color, width, height, htmlCode, helper, profile)
    html_li.inReport = False
    html_li.css({"margin-bottom": '5px'})
    self.val.append(html_li)
    self._lists__map_index.append(category)
    return self

  def __str__(self):
    self._vals = "".join(['''
      <a onclick='this.nextElementSibling.querySelectorAll("li").forEach(
        function(evt){evt.style.display = evt.style.display === "none" ? "" : "none"})' style='cursor:pointer'>%s</a>%s
      ''' % (self._lists__map_index[i] if len(self._lists__map_index) > i else "Category %s" % i, l.html()) for i, l in enumerate(self.val)])
    self.builder_name = self.__class__.__name__
    return "<div %s>%s</div>" % (self.get_attrs(pyClassNames=self.defined), self._vals)






class ListOld(Html.Html):
  cssClsLi = "list-group-item"
  name, category, callFnc = 'Simple List', 'List', 'list'
  _grpCls = CssGrpClsList.CssClassList
  __reqCss, __reqJs = ['bootstrap'], ['jquery', 'bootstrap']
  dashboards = ['DashBoardList']
  cssTitle = "CssTitle4"

  def __init__(self, report, recordSet, icon, title, width, height, draggable, draggableGroupId, draggableMax,
               dataSrc, htmlCode, searchable, selectable, grid, template, globalFilter, profile):
    self.icon, self.title, self.searchable, self.dataSrc, self._jsActions = icon, title, searchable, dataSrc, {}
    self.dsc, self._definedActions = "", []
    if icon is not None:
      self._report.cssImport.add('font-awesome')

    if recordSet is None:
      if dataSrc is None:
        recordSet = []
      else:
        self._report = report
        if dataSrc.get('on_init', False):
          recordSet = self.onInit(None, dataSrc)
    if isinstance(recordSet, dict):
      recordSet = [{"label": k, "value": v} for k, v in recordSet.items()]
    super(List, self).__init__(report, {'vals': recordSet, 'options': {'draggable': draggable, 'groupId': draggableGroupId, 'max': draggableMax}},
                               width=width[0], widthUnit=width[1], height=height[0], heightUnit=height[1],
                               htmlCode=htmlCode, globalFilter=globalFilter, profile=profile)
    self.valFnc = 'GetListData'
    if selectable:
      # Add the selection to a click event
      self._report.style.cssCls('CssBasicListItems', {'cursor': 'pointer'})
      #self.jsFrg('click', "%s['params']['%s'] = [$(this).text()]" % (self._report.jsGlobal.breadCrumVar, self.htmlId))
      #self.valFnc = 'GetListDataSelect'
      #self.addGlobalFnc('GetListDataSelect(htmlId)', "return %(breadCrumVar)s['params'][htmlId];" % {"breadCrumVar": self._report.jsGlobal.breadCrumVar})
    self._jsStyles = {'li': {"padding": "2px 5px 2px 15px"}}
    if not grid:
      self._jsStyles['li']['border'] = 'none'
    self.addGlobalVar("LIST_STATE", "{}")
    if draggable:
      self.css({'cursor': 'pointer', 'background': '#eee'})
    self.css({'overflow': 'auto', 'list-style-type': 'none', 'display': 'inline-block', 'margin': '0 10px 0 0', 'padding': '0 5px 0 5px'})
    if template is not None:
      # Load the list definition from a pre defined template
      # This template can override some existing functions in order to make the list more specific
      # This can be considered (at least I do) as a Mixin !
      import epyk.configs.List

      for script in os.listdir(os.path.dirname(epyk.configs.List.__file__)):
        if script.startswith('Mixin') and script.endswith('py'):
          for name, obj in inspect.getmembers(importlib.import_module("epyk.configs.List.%s" % script.replace(".py", "")), inspect.isclass):
            if hasattr(obj, 'alias') and obj.alias == template:
              self.__class__ = type("%s%s" % (self.__class__.__name__, obj.__name__), (obj, self.__class__), {})
              self.template()
              if self.htmlCode is None: # The htmlId should be reset to the new class name
                self.jsVal = "%s_data" % self.htmlId
              break

          else:
            continue

          break
      else:
        raise Exception("Template configuration %s for list not found !!!" % template)

  @property
  def val(self):
    return "%s('%s')" % (self.valFnc, self.htmlId)

  def onDocumentLoadFnc(self):
    self.addGlobalFnc("%s(htmlObj, data, jsStyles)" % self.__class__.__name__, ''' htmlObj.empty();
      if (Array.isArray(data)) {
        var tmpData = {vals: []};
        data.forEach(function(rec) { 
          if (jsStyles.template != undefined) {tmpData.vals.push(rec)} else {tmpData.vals.push({label: rec})}});
        data = tmpData} 
      else if (Array.isArray( data.vals) && !(typeof data.vals[0] === 'object')) {
        var tmpData = {vals: []};
        data.vals.forEach(function(rec) {tmpData.vals.push( {label: rec})});
        data.vals = tmpData.vals};

      data.vals.forEach(function(rec){
        var label = rec.label; var dataVal = rec.label;
        if (rec.value != undefined) {dataVal = rec.value};
        if (jsStyles.template != undefined) {eval( jsStyles.template )};
        if (jsStyles.events != undefined) {jsStyles.events.forEach( function(rec) { label = label + rec})};
        if (jsStyles.close != undefined) {label = label + jsStyles.close};
        if (rec.badge != undefined){
          var badgeColor = "white";
          if (rec.color != undefined){badgeColor = rec.color};
          label = label + "<span class='badge' style='margin-left:10px;background:%(successColor)s;color:"+ badgeColor +"'>"+ rec.badge +"</span>"}
        if (rec.dsc == undefined) {rec.dsc = ""};
        if (rec.url != undefined) {
          if (rec.target == undefined){rec.target = '_self'};
          label = "<a target='"+ rec.target +"' href='"+ rec.url +"' style='color:%(blackColor)s'>" + label + "</a>"}
        else if(rec.script != undefined){
          var paramLinks = []; for(var k in rec.params){paramLinks.push(k +"="+ rec.params[k])};
          if (rec.target == undefined){rec.target = '_self'};
          label = "<a target='"+ rec.target +"' href='%(url)s/run/%(report_name)s/"+ rec.script +"?"+ paramLinks.join("&") +"' style='color:%(blackColor)s'>" + label + "</a>"
        };
        var li = $('<li data-placement="right" data-value="'+ dataVal +'" title="'+ rec.dsc +'" class="%(itemStyle)s"></li>').css(jsStyles['li']);
        li.append(toMarkup(label+''));
        htmlObj.append(li)
      }); 

      htmlObj.find('li').tooltip(); 
      if ( data.options != undefined && data.options.draggable){
         if ( data.options.groupId != null) {
          htmlObj.addClass( data.options.groupId ) ; 
          if ( data.options.max != null ) { 
            htmlObj.sortable( {placeholder: "ui-state-highlight", 
              receive: function(event, ui) { if (htmlObj.children().length > data.options.max) { $(ui.sender).sortable('cancel'); } },
              dropOnEmpty: true, connectWith: "." + data.options.groupId} )
          } else {
            htmlObj.sortable( {placeholder: "ui-state-highlight", dropOnEmpty: true, connectWith: "." + data.options.groupId} ) 
          }}
         else { 
          if ( data.options.max != null ) { 
            htmlObj.sortable( {placeholder: "ui-state-highlight", 
              receive: function(event, ui) { if (htmlObj.children().length > data.options.max) { $(ui.sender).sortable('cancel'); } },
              dropOnEmpty: true, connectWith: "." + data.options.groupId} )
          } else {
            htmlObj.sortable( {placeholder: "ui-state-highlight", dropOnEmpty: true, connectWith: "." + data.options.groupId} ) 
          } } ;
         htmlObj.disableSelection() ;
      } ''' % {"blackColor": self.getColor('greys', -1), "url": self._report._urlsApp['report'], 'report_name': self._report.run.report_name,
               "successColor": self.getColor('success', 1), "itemStyle": self._report.style.cssName('CssBasicListItems')})

  def __add__(self, rec):
    """ Add items to a container """
    self.vals['vals'].append(rec)
    return self

  @property
  def eventId(self): return "#%s li" % self.htmlId

  @property
  def jsQueryData(self):
    return "{event_val: $(event.currentTarget).data('value'), event_label: $(event.currentTarget).text(), event_data: %s, event_code: '%s' }" % (self.val, self.htmlId)

  def jsLoadFromSrc(self, jsData='data', jsDataKey=None, isPyData=False):
    if self.dataSrc['type'] == 'flask':
      url = url_for("%s.%s" % (self.dataSrc['blueprint'], self.dataSrc['fnc'].__name__), **dict(zip(self.dataSrc['pmts_def'], self.dataSrc['pmts'])) )
      return self._report.jsPost(url, jsData=self.dataSrc.get('htmlObjs'), htmlCodes=self.dataSrc.get('htmlCodes'), jsFnc=[self.jsGenerate(jsData, jsDataKey=jsDataKey, isPyData=isPyData)])
    else:
      return self._report.jsPost(self.dataSrc['script'], jsData=self.dataSrc.get('htmlObjs'), htmlCodes=self.dataSrc.get('htmlCodes'), jsFnc=[self.jsGenerate(jsData, jsDataKey=jsDataKey, isPyData=isPyData)] )

  def jsDisable(self, jsData='data', jsDataKey=None, isPyData=False, jsParse=False):
    jsData = self._jsData(jsData, jsDataKey, jsParse, isPyData)
    return '''
      const ldata = %(jsData)s;
      %(jqId)s.find('li').removeClass("%(disableCls)s");
      %(jqId)s.find('li').each(function(){
        if (ldata.indexOf($(this).text()) > -1){$(this).addClass("%(disableCls)s")}
      })''' % {'jsData': jsData, 'jqId': self.jqId, 'disableCls': self._report.style.cssName("CssBasicListItemsDisabled")}

  def jsHideItems(self, jsData='data', jsDataKey=None, isPyData=False, jsParse=False):
    jsData = self._jsData(jsData, jsDataKey, jsParse, isPyData)
    return '''
      var ldata = %(jsData)s;
      %(jqId)s.find('li').each(function(){
        if(ldata == ''){$(this).show()}
        else if ($(this).text().toLowerCase().indexOf(ldata.toLowerCase()) == -1){$(this).hide()}
        else {$(this).show()}
      })''' % {'jsData': jsData, 'jqId': self.jqId}

  def click(self, jsFncs):
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    self._report.cssUpdate('CssBasicListItems', 'cursor', 'pointer')
    return self.jsFrg('click', ";".join(jsFncs) if isinstance(jsFncs, list) else jsFncs)

  def delete(self, url=None, jsData=None, jsFncs='', cacheObj=None, isPyData=True, isDynUrl=False, httpCodes=None, htmlCodes=None):
    """
    Set a delete button to each item in the list. This can be enhanced with a special Ajax or simple javascript function

    Example
    myObj.delete( jsFncs=[report.jsPost("test.py", jsFnc=[report.jsConsole() ]) ])
    """
    self._jsStyles['close'] = "<i onclick='DeleteItem(event, this)' style='position:absolute;right:5px;color:#C00000' class='far fa-times-circle'></i>"
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    jsFncs.append("$(srcObj).parent().remove()")
    if url is not None:
      jsFncs = self._report.jsPost(url=url, jsData=jsData, jsFncs=jsFncs, cacheObj=cacheObj, isPyData=isPyData, isDynUrl=isDynUrl,
                                   httpCodes=httpCodes, htmlCodes=htmlCodes)
    self.addGlobalFnc('DeleteItem(event, srcObj)', '''
          var data = $(srcObj).parent().text() ;
          %(jsFncs)s; event.stopPropagation(); ''' % {'jsFncs': ';'.join(jsFncs)}, "Delete an item from the list")
    return self

  def jsAction(self, action, icon, pyCssCls="CssSmallIcon", tooltip="", url=None, jsData=None, jsFncs=None, httpCodes=None):
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs] if jsFncs is not None else []

    # Add this to an ajax POST call if an URL is defined
    fnc = self._report.jsPost(url=url, jsData=jsData, jsFnc=jsFncs, httpCodes=httpCodes) if url is not None else ";".join(jsFncs)
    self._report.style.cssCls(pyCssCls)
    self._jsActions[action] = "<span id='%(htmlId)s_%(action)s' title='%(tooltip)s' class='%(cssStyle)s %(icon)s'></span>" % {
      "icon": icon, "cssStyle": self._report.style.cssName(pyCssCls), "htmlId": self.htmlId, 'tooltip': tooltip, 'action': action}
    self._report.jsOnLoadFnc.add("$('#%(htmlId)s_%(action)s').on('click', function(event) { %(jsFncs)s; })" % {"htmlId": self.htmlId, "jsFncs": fnc, 'action': action})
    if action not in self._definedActions:
      self._definedActions.append(action)
    return self

  def jsItemAction(self, icon, jsFncName, jsFncs):
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    if not 'events' in self._jsStyles:
      self._jsStyles['events'] = []
    self.css({'cursor': 'pointer'})
    self._jsStyles['events'].append("<i name='%s' style='position:absolute;right:%spx;color:black' class='%s'></i>" % (jsFncName, (len( self._jsStyles['events'])+1) * 20 + 5, icon ))
    self._report.jsOnLoadFnc.add('''$(document).on('click', 'i[name="%s"]' , function(event) { var data = $(this).parent().text() ; %s })''' % (jsFncName, ';'.join(jsFncs)) )

  def jsAdd(self, jsData='data', jsDataKey=None, isPyData=False, jsParse=False):
    """
    This function will allow to set a function to add from the browser and the user input new entries in a list
    """
    data = self._jsData(jsData, jsDataKey, jsParse, isPyData)
    return '''
      %(jqId)s.append('<li class="%(clsLi)s" style="white-space:nowrap;padding:5px;background-color:%(lightBlue)s">' + %(data)s + '<i style="float:right;cursor:pointer" onclick="$(this).parent().remove()" class="far fa-times-circle"></i></li>') ;
      ''' % {'clsLi': self.cssClsLi, 'data': data, 'jqId': self.jqId, 'lightBlue': self.getColor('colors', 4)}

  def jsEvents(self):
    if hasattr(self, 'jsFncFrag'):
      for eventKey, fnc in self.jsFncFrag.items():
        if self.htmlCode is not None:
          fnc.insert(0, self.jsAddUrlParam(self.htmlCode, self.val, isPyData=False))
        if getattr(self._report, 'PROFILE', False):
          self._report.jsOnLoadEvtsFnc.add('''
            $(document).on('%(eventKey)s', '%(eventId)s', function(event){var t0_event = performance.now(); 
                if (!$(this).hasClass('%(disableCls)s')){
                  var useAsync = false; var data = %(data)s; %(jsInfo)s; %(jsFnc)s; 
                  if (!useAsync) {
                    var body_loading_count = parseInt($('#body_loading span').text()); $('#body_loading span').html(body_loading_count - 1);
                    if ($('#body_loading span').html() == '0') { $('#body_loading').remove()}};
                console.log('|LIst|%(eventKey)s(%(htmlId)s)|'+ (performance.now()-t0_event))
            }})''' % {'eventId': self.eventId, 'eventKey': eventKey, 'data': self.jsQueryData, 'disableCls': self._report.style.cssName("CssBasicListItemsDisabled"),
                      'htmlId': self.htmlId, 'jsFnc': ";".join([f for f in fnc if f is not None]), 'jsInfo': self._report.jsInfo('process(es) running', 'body_loading')})
        else:
          self._report.jsOnLoadEvtsFnc.add('''
            $(document).on('%(eventKey)s', '%(eventId)s', function(event){
                if (!$(this).hasClass('%(disableCls)s')){
                  var useAsync = false; var data = %(data)s; %(jsInfo)s; %(jsFnc)s; 
                  if (!useAsync) {
                    var body_loading_count = parseInt($('#body_loading span').text()); $('#body_loading span').html(body_loading_count - 1);
                    if ($('#body_loading span').html() == '0') { $('#body_loading').remove()}}
            }})''' % {'eventId': self.eventId, 'eventKey': eventKey, 'data': self.jsQueryData, 'disableCls': self._report.style.cssName("CssBasicListItemsDisabled"),
                      'jsFnc': ";".join([f for f in fnc if f is not None]), 'jsInfo': self._report.jsInfo('process(es) running', 'body_loading')})

  # -----------------------------------------------------------------------------------------
  #                                    LIST EXPORT OPTIONS
  # -----------------------------------------------------------------------------------------
  def __str__(self):
    self.addGlobalFnc('GetListData(htmlId)', '''
          var res = [];
          $('#'+ htmlId + ' li').each(function(item){
            if ( $(this).children().length == 1) {res.push($(this).find('a').html())}
            else { res.push($( this ).html())}}); return res''')
    self.addGlobalFnc('SearchListData(htmlId)', '''
      $('#'+ htmlId + ' li').each( function(item){
        if ($(this).text().indexOf( $('#'+ htmlId + '_search').val()) > -1) {$(this).show()}
        else { $(this).hide(); $('#'+ htmlId).css({'overflow': 'hidden'})}
      })''')
    icon = '<div style="width:100%%;text-align:center;font-size:50px">' % self.icon if self.icon is not None else ''
    searchableStr, events = '', []
    for action in self._definedActions:
      if action in self._jsActions:
        events.append(self._jsActions[action])

    css_mod, title4 = self._report.style.get(self.cssTitle), ""
    if css_mod is not None:
      self._report.style.cssCls(self.cssTitle)
      title4 = css_mod.classname
    if self.dsc != '':
      self.dsc = "<div style='width:100%%;padding-bottom:5px;text-align:justify'>%s</div>" % self.dsc
    if self.searchable:
      searchableStr = '''
        <div style="width:100%%;padding:5px 0 5px 0;display:inline-block">
          <div onclick="SearchListData('%(htmlId)s')" style="margin-bottom:-60px;margin-right:4px;position:relative;float:right;cursor:hand;cursor:pointer;height:29px;font-size:12px;"><i class="fas fa-search"></i></div>
          <input onkeyup="SearchListData('%(htmlId)s')" id="%(htmlId)s_search" style="width:100%%;color:#777;" type="text" placeholder="Search for" />
        </div>
        ''' % {'htmlId': self.htmlId}
    return '''
      <div %(attr)s>
        <div style="margin-bottom:5px">
          <i class="%(icon)s"></i><span class="%(title4)s">%(title)s</span>
           %(events)s
        </div>
        %(searchableStr)s
        %(dsc)s
        <ul id="%(htmlId)s" class="list-group"></ul>
      </div>%(helper)s''' % {'searchableStr': searchableStr, 'htmlId': self.htmlId, "attr": self.get_attrs(pyClassNames=self.defined, withId=False), 'blackColor': self.getColor('greys', 9),
                             "icon": icon, "title": self.title, "helper": self.helper, "events": "".join(events),
                             "title4": title4, "dsc": self.dsc}

  def to_word(self, document):
    from docx.shared import RGBColor

    for rec in self.vals['vals']:
      if isinstance(rec, str):
        p = document.add_paragraph(style='ListBullet')
        runner = p.add_run(rec)
        if self._report.http.get(self.htmlCode) == rec:
          runner.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
      else:
        document.add_paragraph(rec['label'], style='ListBullet')
        runner = p.add_run(rec)
        if self._report.http.get(self.htmlCode) == rec['label']:
          runner.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

  def to_xls(self, workbook, worksheet, cursor):
    if self.title != '':
      cell_format = workbook.add_format({'bold': True})
      worksheet.write(cursor['row'], 0, self.title, cell_format)
      cursor['row'] += 1
    for rec in self.vals['vals']:
      cell = rec if isinstance(rec, str) else rec['label']
      cell_format = workbook.add_format({})
      if self._report.http.get(self.htmlCode) == cell:
        cell_format = workbook.add_format({'bold': True, 'font_color': self.getColor('colors', 5)})
      worksheet.write(cursor['row'], 0, cell, cell_format)
      cursor['row'] +=1
    cursor['row'] += 1


  # --------------------------------------------------------------------------------------------------------------
  #                                   TEMPLATE ITEM SECTION
  #
  # Simple templates should be defined in the config / List section
  # The ones below are more complex and they are using some events
  #
  def _template(self, strTemplate):
    """
    Python template to display the content of an item in the list.
    This should be a string using the Python formatting

    Example
    myObj.template("%(filename)s, last update %(lst_update)s")
    """
    strTemplate = "".join([line.strip() for line in strTemplate.split('\n')]).replace('"', "'") # Cannot use the " in the string as this is the javascript delimiter
    formatStr = re.compile("%\(([0-9a-zA-Z_]*)\)s")
    matches = formatStr.findall(strTemplate)
    if matches:
      for res in formatStr.finditer(strTemplate):
        strTemplate = strTemplate.replace(res.group(0), '"+ rec["%s"] +"' % res.group(1))
    self._jsStyles['template'] = 'value = "%s"' % strTemplate
    return self

  def templateMessageDetails(self, withHr=True, color='black', icon=''):
    """
    Python template to display the content of an message item.
    The record expected for those items are a bit special.
    Each items should get the below keys

    Example
    myObj.templateMessageDetails(withHr=True) # If bottom border
    """
    if icon != '':
      icon = '<i class="%s" style="margin-right:5px"></i>' % icon
    self._template('''
      <div style="width:100%;display:inline-block">
        <div style="width:60px;float:left">
          <div style="font-size:14px">%(views)s <span style="font-size:10px">views</span></div>
          <div style="font-size:14px">%(answers)s <span style="font-size:10px">answers</span></div>
          <div style="font-size:14px">%(interest)s <span style="font-size:10px">interest</span></div>
        </div>
        <div>
          <a href='%(url)s' target='_blank' style='color:'''+ color +''';text-decoration:none;font-weight:bold;font-size:14px;color:"+ self.getColor('colors', 8) +"'>'''+ icon +'''%(title)s</a>
          <div style="white-space:pre-wrap;"></div>
        </div>
      </div>
      ''' + '<hr />' if withHr else '' )
    return self

  def templateIdeas(self, jsFncs, icon="fas fa-users", tooltip="Support Idea"):
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    self.addGlobalFnc('IdeaEvent(scrId)', ''' var data = {event_val: scrId } ; %(jsFncs)s ''' % {"jsFncs": ";".join(jsFncs)})

    self._template('''
      <div style="width:100%;border:1px solid %(color)s;border-radius:2px;padding:5px 10px 5px 10px;margin-bottom:5px">
        <span style="font-weight:bold;font-size:14px">%(title)s</span>
        <i onclick="IdeaEvent(%(id)s)" class="''' + icon + '''" title="''' + tooltip + '''" style="margin-left:15px;cursor:pointer"><span class="badge" style="color:black">%(count)s</span></i>
        <br />
        <div style="width:100%;display:inline-block;margin-top:5px">
          <div style="display:inline-block;float:left;width:50px;text-align:center;">
            <i class="%(icon)s" style="color:%(color)s;font-size:20px;margin-left:5px"></i>
          </div>
          <div style="white-space:pre-wrap;">%(content)s</div>
        </div>
      </div>''' )
    return self

  def templateTags(self, url=None, jsFncs='', cacheObj=None, isDynUrl=False, httpCodes=None, htmlCodes=None):
    """
    Python template to display the content of an item with a badge value.
    The record with this template should follow the below structure:
    record = { "value": "", "id": ""}
    """
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    jsFncs.append("$(srcObj).remove() ")
    if url is not None:
      js = [self._report.jsPost(url, jsFnc=jsFncs, cacheObj=cacheObj, isDynUrl=isDynUrl, httpCodes=httpCodes, htmlCodes=htmlCodes)]
    else:
      js = jsFncs
    self.addGlobalFnc('RemoveListTag(srcObj)', "var data = { event_val: $(srcObj).text() };  %s" % ";".join(js))
    self._template('''
      <div style="width:100%;border-radius:2px;padding:1px 10px 1px 10px;margin-bottom:5px">
        <span style="color:black;background-color:'''+ self.getColor('colors', 2) +''';border-radius:5px;padding:4px 10px;">%(value)s<i onclick="RemoveListTag($(this).parent())" class="fas fa-times" style="margin-left:20px;cursor:pointer"></i></span>
      </div>''' )
    return self

  def templateUrl(self, target='_self', color='black'):
    """
    Python template to display the content of an item with URL link
    The record expected for those items are a bit special.
    Each items should get the below keys
        record = { "url": "", "title": ""}
    """
    self._template(''' <a href="%(url)s" style="white-space:pre-wrap;color:'''+ color +'''" target="'''+ target +'''">%(title)s</a>''')
    return self

  def templateAnswer(self, url=None, jsFncs='', cacheObj=None, isDynUrl=False, httpCodes=None, htmlCodes=None):
    """
    Python template to display the items corresponding to an answer. Those items are visible in the section Questions / details in the framework.
    TThe record expected for those items are a bit special.
    Each items should get the below keys
        record = {"title": "", "borderColor": "", "color": "", "id": "", "icon": "", "author": "", "interest": "", "lst_mod_dt": "" }
    """
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    jsFncs.append("srcObj.append( '<span style=\"font-style:italic;width:100%;display:block\">' + msg + '</span>' )")
    if url is not None:
      js = [self._report.jsPost(url, jsFnc=jsFncs, cacheObj=cacheObj, isDynUrl=isDynUrl, httpCodes=httpCodes, htmlCodes=htmlCodes)]
    else:
      js = jsFncs
    self.addGlobalFnc('AddExtraComment(srcObj, msg, rowId)', "var data = { event_val: msg, event_id: rowId }; %s" % ";".join(js))
    self.addGlobalFnc('AnswerInterest(answerId, value)', self._report.jsPost("'/questions/answer/ranking/' + answerId + '/' + value", isDynUrl=True))

    self._template('''
      <div style="display:block;width:100%;border-radius:5px;border:1px solid %(borderColor)s;background-color:%(borderColor)s;color:%(color)s">
        <div style="display:block;width:50px;float:left;height:100%;">
          <i onclick="AnswerInterest(%(id)s, 1)" class="fas fa-caret-up" style="cursor:pointer;color:#C9CBCF;width:100%%;display:block;font-size:40px;text-align:center"></i>
          <span style="color:#C9CBCF;width:100%%;display:block;font-size:30px;text-align:center"> %(interest)s </span>
          <i onclick="AnswerInterest(%(id)s, -1)" class="fas fa-caret-down" style="cursor:pointer;color:#C9CBCF;width:100%%;display:block;font-size:40px;text-align:center"></i>
        </div>
        <div style='padding:10px;text-align:left;display:block;margin-left:50px;'>
          <span style='font-size:14px;display:inline-block;width:100%;white-space:pre-wrap;line-height:1.4;'>%(answer)s</span>
          <div style='margin-top:10px;text-align:right;width:100%;font-style:italic;'>%(author)s, %(lst_mod_dt)s</div>
          <hr style="margin-bottom:0px">
          <div id='%(id)s_extra' style='margin:0px;padding:0px'>%(xtra)s</div>
          <input type='text' class='form-control' style='margin-top:5px;width:100%' onkeydown='if (event.keyCode == 13) {  AddExtraComment($(this).prev(), $(this).val(), %(id)s ) }' />
          <hr style="margin:10px 20%;background:#888888">
        </div> 
      </div> ''')

    self._jsStyles['template'] = 'var xtra = []; rec["xtra"].forEach( function(val) { xtra.push("<span style=\'white-space:pre-wrap;font-style:italic;width:100%%;display:block\'>" + val + "</span>"); } ); rec["xtra"]=xtra.join(""); %s;' % self._jsStyles['template']
    return self


  # -----------------------------------------------------------------------------------------
  #                                    MARKDOWN SECTION
  # -----------------------------------------------------------------------------------------
  @staticmethod
  def matchMarkDownBlock(data): return re.match(">>>List", data[0])

  @staticmethod
  def matchEndBlock(data): return data.endswith("<<<")

  @classmethod
  def convertMarkDownBlock(cls, data, report=None):
    records, pmts, css = [], {'selectable': False}, {}
    for val in data[1:-1]:
      val = val.strip()
      if val.startswith("@"):
        dataAttr = val[1:].strip().split(";")
        for d in dataAttr:
          a, b = d.split(":")
          css[a] = b
      elif val.startswith("--"):
        dataAttr = val[2:].strip().split(";")
        for d in dataAttr:
          a, b = d.split(":")
          pmts[a] = b
      else:
        records.append({"label": val.strip()})
    if report is not None:
      getattr(report, 'list')(records, **pmts).css(css)
    return ["report.list(%s, selectable=False)" % json.dumps(records)]

  @classmethod
  def jsMarkDown(self, vals):
    return [">>>List", [rec for rec in vals['label']], "<<<"]


class ListBadge(Html.Html):
  cssCls, cssClsLi = ['list-group'], "list-group-item"
  _grpCls = CssGrpCls.CssGrpClassBase
  __reqCss, __reqJs = ['bootstrap'], ['bootstrap']
  name, category, callFnc = 'List Badges', 'Container', 'listbadge'

  def __init__(self, report, recordSet, color, size, width, height, draggable, draggableGroupId,
               draggableMax, dataSrc, profile):
    # if recordSet:
    #   if isinstance(recordSet, list) and isinstance(recordSet[0], str):
    #     recordSet = [{'label': val, 'value': 0} for val in recordSet]
    super(ListBadge, self).__init__(report, {'vals': recordSet, 'options': {'draggable': draggable, 'groupId': draggableGroupId, 'max': draggableMax}},
                                    width=width[0], widthUnit=width[1], height=height[0], heightUnit=height[1], dataSrc=dataSrc, profile=profile)
    self.size = "%s%s" % (size[0], size[1])
    self.css({"font-size": self.size})
    self.val['color'] = self.getColor('success', 1) if color is None else color

  def onDocumentLoadFnc(self):
    self.addGlobalFnc("%s(htmlObj, data, jsStyles)" % self.__class__.__name__, ''' htmlObj.empty();
      if (data.vals == undefined){data = {vals: data, options: {}}};
      data.vals.forEach(function(rec){
        if (rec.url != undefined) {var content = '<a href="'+ rec.url +'">'+ rec.label +'</a>'} else {var content = rec.label};
        if (rec.color == undefined) {rec.color = '%(color)s'};
        htmlObj.append('<li class="%(clsLi)s">'+ content +'&nbsp;<span class="badge" style="background-color:'+ rec.color +';color:white">'+ rec.value +'</span></li>'); 
      });
      if (data.options.draggable){
         if (data.options.groupId != null) {
          htmlObj.addClass(data.options.groupId); 
          if (data.options.max != null) { 
            htmlObj.sortable( {placeholder: "ui-state-highlight", 
              receive: function(event, ui) {if (htmlObj.children().length > data.options.max) { $(ui.sender).sortable('cancel')}},
              dropOnEmpty: true, connectWith: "."+ data.options.groupId})
          } else {
            htmlObj.sortable({placeholder: "ui-state-highlight", dropOnEmpty: true, connectWith: "." + data.options.groupId})
          }}
         else { 
          if (data.options.max != null) { 
            htmlObj.sortable( {placeholder: "ui-state-highlight", 
              receive: function(event, ui) {if (htmlObj.children().length > data.options.max) { $(ui.sender).sortable('cancel')}},
                dropOnEmpty: true, connectWith: "."+ data.options.groupId})
          } else {
            htmlObj.sortable({placeholder: "ui-state-highlight", dropOnEmpty: true, connectWith: "."+ data.options.groupId})}};
         htmlObj.disableSelection()
      }''' % {'color': self.getColor('success', 1), 'clsLi': self.cssClsLi}, 'Javascript Object builder')

  def __str__(self):
    # self.addGlobalFnc('GetListData(htmlId)', '''
    #   var res = []; $('#'+ htmlId + ' li').each(function(item){
    #   if ($(this).children().length == 1) { res.push($( this ).find('a').html()) ;} else { res.push($( this ).html()) ;}}); return res;''',
    #   'Returns the HTML text of all the items in the object')
    return '<div %s><ul></ul></div>%s' % (self.get_attrs(pyClassNames=self.defined), self.helper)

  @staticmethod
  def matchMarkDownBlock(data): return re.match(">>>Badge", data[0])

  @staticmethod
  def matchEndBlock(data): return data.endswith("<<<")

  @classmethod
  def convertMarkDownBlock(cls, data, report):
    recordSet = []
    for val in data[1:-1]:
      splitData = val.split("|")
      if len(splitData) == 1:
        splitData.append(0)
      recordSet.append({'label': splitData[0], "value": splitData[1]})
    if report is not None:
      getattr(report, 'listbadge')(recordSet)
    return ["report.listbadge(%s)" % json.dumps(recordSet)]

  @classmethod
  def jsMarkDown(self, vals):
    return [">>>Badge", ["%s|%s" % (rec['value'], rec['label']) for rec in vals['vals']], "<<<"]


class HtmlListAccordeon(Html.Html):
  name, category, callFnc = 'Vertical expendable list', 'List', 'accordeon'
  __reqCss, __reqJs = ['bootstrap', 'font-awesome'], ['bootstrap', 'font-awesome']
  _grpCls = CssGrpClsList.CssClassListAccordeon

  def __init__(self, report, recordSet, color, width, size, dataSrc, profile):
    super(HtmlListAccordeon, self).__init__(report, recordSet, width=width[0], widthUnit=width[1], dataSrc=dataSrc, profile=profile)
    size = "%s%s" % (size[0], size[1])
    self.css({'color': color if color is not None else 'inherit', 'font-size': size})

  def onDocumentLoadFnc(self):
    self.addGlobalFnc("%s(htmlObj, data)" % self.__class__.__name__, '''htmlObj.empty();
      var categories = {}; var cats = []; var catsIcons = {};  
      data.forEach(function(rec){
        if(rec.icon != undefined){catsIcons[rec.category] = '<i class="'+ rec.icon +'"></i>&nbsp;'};
        if(rec.category in categories){categories[rec.category].push(rec)} 
        else{categories[rec.category] = [rec]; cats.push(rec.category)}});
      htmlId = htmlObj.attr('id'); var countItems = 0;   
      cats.forEach(function(cat){
        if(cat in catsIcons){var liItem = $('<li %(cssLi)s name="'+ htmlId +'_menu" id="'+ htmlId +'_menu_'+ countItems +'"><a href="#">'+ catsIcons[cat] + cat +'</a></li>')}
        else {var liItem = $('<li %(cssLi)s name="'+ htmlId +'_menu" id="'+ htmlId +'_menu_'+ countItems +'"><a href="#'+ htmlId +'_menu_'+ countItems +'">'+ cat +'</a></li>')}
        countItems = countItems + 1; var ulItem = $('<ul></ul>');
        categories[cat].forEach(function(rec){
          if(rec.color != undefined){var content = 'style="color:'+ rec.color +'"'} else{var content = ''};
          ulItem.append('<li %(cssItems)s '+ content +'><a href="'+ rec.url +'">'+ rec.value +'</a></li>')
        }); liItem.append(ulItem); htmlObj.append(liItem)
      })''' % {'cssLi': self._report.style.getClsTag(['CssListNoDecoration', 'CssHreftMenu', 'CssListLiUlContainer'], loadCls=True),
               'cssItems': self._report.style.getClsTag(['CssListNoDecoration', 'CssHrefSubMenu', 'CssListLiSubItem'], loadCls=True)}, 'Javascript Object builder')

  def __str__(self):
    self._report.jsOnLoadFnc.add("$('li[name=%(htmlId)s_menu]').click(function(){$('#'+ this.id +' ul').toggle()})" % {'htmlId': self.htmlId})
    return '<div %s></div>' % self.get_attrs(pyClassNames=self.defined)


class CheckList(Html.Html):
  name, category, callFnc = 'List Checked', 'List', 'checklist'
  _grpCls = CssGrpCls.CssGrpClassBase

  def __init__(self, report, recordSet, width, height, globalFilter, dataSrc, profile):
    if isinstance(recordSet[0], str):
      recordSet = [{'value': r} for r in recordSet]
    super(CheckList, self).__init__(report, recordSet, width=width[0], widthUnit=width[1], height=height[0],
                                    heightUnit=height[1], globalFilter=globalFilter, dataSrc=dataSrc, profile=profile)
    self.css({"padding-left": "20px"})

  @property
  def jqId(self): return "$('#%s')" % self.htmlId

  def onDocumentLoadFnc(self):
    self.addGlobalFnc("%s(htmlObj, data, jsStyles)" % self.__class__.__name__, ''' htmlObj.empty() ;
        var cat = htmlObj.attr('id') +"_cat" ;
        data.forEach(function(rec){
          if (rec.isChecked == undefined) { 
            if (rec.disabled != undefined) { htmlObj.append('display:inline-flex;align-items:center"><input style="margin-right:5px" type="radio" name="'+ cat + '" disabled value="' + rec.value + '">' + rec.value + '</label><br />') ; }
            else {htmlObj.append('<label style="display:inline-flex;align-items:center"><input style="margin-right:5px" type="radio" name="'+ cat + '" value="' + rec.value + '">' + rec.value + '</label><br />')} } 
          else {htmlObj.append('<label style="display:inline-flex;align-items:center"><input style="margin-right:5px" type="radio" name="'+ cat + '" value="' + rec.value + ' " checked>'+ rec.value + '</label><br>')}
        })''', 'Javascript Object builder')

  def __str__(self):
    return '<div %s></div>%s' % (self.get_attrs(pyClassNames=self.defined), self.helper)

  @classmethod
  def matchMarkDownBlock(cls, data): return re.match(">>>%s" % cls.callFnc, data[0])

  @staticmethod
  def matchEndBlock(data): return data.endswith("<<<")

  @classmethod
  def convertMarkDownBlock(cls, data, report=None):
    recordSet = []
    for val in data[1:-1]:
      line = val.strip()
      if line.startswith("- [x] "):
        recordSet.append({'value': line[5:], "isChecked": 1})
      else:
        recordSet.append({'value': line[5:], 'disabled': 1})
    if report is not None:
      getattr(report, 'checklist')(recordSet)
    return ["report.checklist(%s)" % json.dumps(recordSet)]

  @classmethod
  def jsMarkDown(cls, vals):
    result = []
    for rec in vals:
      if rec.get('isChecked') == 1:
        result.append("- [X] %s" % rec['value'])
      else:
        result.append("- [ ] %s" % rec['value'])
    return [">>>%s" % cls.callFnc, result, "<<<"]


class ListTournaments(Html.Html):
  name, category, callFnc = 'Brackets', 'Container', 'brackets'
  __reqCss, __reqJs = ['jquery-brackets'], ['jquery-brackets']

  def __init__(self, report, records, width, height, options, profile):
    self.options = {} if options is None else options
    super(ListTournaments, self).__init__(report, {'vals': records, 'save': 'null', 'edit': 'null', 'render': 'null',
                                                   'options': self.options}, width=width[0], widthUnit=width[1],
                                                   height=height[0], heightUnit=height[1], profile=profile)
    self.css({'overflow': 'auto', "padding": "auto", "margin": "auto"})

  def addFnc(self, fncName, jsFncs):
    if isinstance(jsFncs, list):
      jsFncs = ";".join(jsFncs)
    self.vals[fncName] = jsFncs

  def onDocumentLoadFnc(self):
    # , disableToolbar: true, disableTeamEdit: false
    self.addGlobalFnc("%s(htmlObj, data)" % self.__class__.__name__, ''' htmlObj.empty() ;
      parameters = {centerConnectors: true, init: data.vals }; 
      if (data.save != "null"){parameters['save'] = new Function('rec', 'userData', 'var data = {challenge: JSON.stringify(rec), userProno: JSON.stringify(userData) } ;' + data.save) };
      if (data.render != "null"){parameters['decorator'] = {render: new Function('rec', 'userData', data.save), edit: function(container, data, doneCb) { } } };
      if (data.edit != "null"){ 
        if ( data.render == "null" ) { parameters['decorator']['render'] = function(rec, userData) {} } ;
        parameters['decorator']['edit'] = new Function('container', 'data', 'doneCb', data.edit); };
      for (var k in data.options) { parameters[k] = data.options[k] ;};
      htmlObj.bracket( parameters )''', 'Javascript Object builder')

  def __str__(self):
    return "<div %s></div>" % self.get_attrs(pyClassNames=self.defined)
