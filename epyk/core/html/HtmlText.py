
import re
import os
import json

from epyk.core.html import Html
from epyk.core.html.options import OptText

from epyk.core.html import Defaults as Default_html

# The list of Javascript classes
from epyk.core.js.objects import JsNodeDom
from epyk.core.js.html import JsHtml
from epyk.core.css.styles import GrpCls


class Label(Html.Html):
  name, category, callFnc = 'Label', 'Text', 'label'

  def __init__(self, report, text, color, align, width, height, htmlCode, tooltip, profile, dflt_options):
    super(Label, self).__init__(report, text, css_attrs={"width": width, "height": height, 'color': color, 'text-align': align},
                                code=htmlCode, profile=profile)
    self.css({'margin': '0 5px', 'float': 'left', 'display': 'inline-block', 'line-height': '23px',
              'vertical-align': 'middle', 'text-align': 'left'})
    self._jsStyles = dflt_options
    if tooltip:
      self.set_attrs(name='title', value=tooltip)

  @property
  def id_html(self):
    """

    Documentation
    https://developer.mozilla.org/fr/docs/Web/API/Element/getElementsByTagName

    :return:
    """
    return JsNodeDom.JsDoms.get("document.getElementById('%s')" % self.htmlId)

  def click(self, jsFncs, profile=False):
    """
    Add a click event for a component

    The event will be automatically added to the onload section to be activated once the component
    has been build

    Example
    select.label.click(str(rptObj.js.console.log("test")))

    Documentation
    https://www.w3schools.com/js/js_htmldom_eventlistener.asp
    https://www.w3schools.com/jsref/event_onload.asp

    :param jsFncs: An array of Js functions or string. Or a string with the Js
    :param profile:

    :return: The htmlObj
    """
    self.css({"cursor": "pointer"})
    self.on("click", jsFncs)
    return self

  def selectable(self, flag=False):
    """
    Description:
    ------------
    Make the label component not selectable.

    This will be done by adding the class CssTextNotSelectable to the component

    Attributes:
    ----------
    :param flag: Boolean.

    :return: self to allow the chains
    """
    if not flag:
      self.style.add_classes.text.no_selection()
    return self

  @property
  def _js__builder__(self):
    mark_up = self._report.js.string("data", isPyData=False).toStringMarkup()
    return '''
      var content = options.markdown > 0 ? %(markUp)s : data
      if(options._children > 0){htmlObj.insertAdjacentHTML('beforeend', '<div style="display:inline-block;vertical-align:middle">'+ content +'</div>')}
      else{htmlObj.innerHTML = content};
      if(typeof options.css !== 'undefined'){for(var k in options.css){htmlObj.style[k] = options.css[k]}}''' % {"markUp": mark_up}
    
  def __str__(self):
    return '<label %s>%s</label>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), self.val, self.helper)


class Span(Html.Html):
  name, category, callFnc = 'Span', 'Texts', 'span'

  def __init__(self, report, text, color, align, width, height, htmlCode, tooltip, profile):
    super(Span, self).__init__(report, text, css_attrs={"width": width, "height": height, "color": color, 'text-align': align},
                               code=htmlCode, profile=profile)
    self.css({'line-height': '%spx' % Default_html.LINE_HEIGHT, 'margin': '0 5px', 'display': 'inline-block', 'vertical-align': 'middle'})
    if tooltip is not None:
      self.tooltip(tooltip)

  @property
  def id_container(self):
    return self.htmlId

  @property
  def id_jquery(self):
    return JsNodeDom.JsDoms.get("$('#%s')" % self.htmlId)

  @property
  def id_html(self):
    """

    Documentation
    https://developer.mozilla.org/fr/docs/Web/API/Element/getElementsByTagName

    :return:
    """
    return JsNodeDom.JsDoms.get("document.getElementById('%s')" % self.htmlId)

  @property
  def dom(self):
    """
    Javascript Functions

    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object

    :rtype: JsHtml.JsHtmlRich
    """
    if self._dom is None:
      self._dom = JsHtml.JsHtmlRich(self, report=self._report)
    return self._dom

  def click(self, jsFncs, profile=False):
    """
    Add a click event for a component

    The event will be automatically added to the onload section to be activated once the component
    has been build

    Example
    select.label.click(str(rptObj.js.console.log("test")))

    Documentation
    https://www.w3schools.com/js/js_htmldom_eventlistener.asp
    https://www.w3schools.com/jsref/event_onload.asp

    :param jsFncs: An array of Js functions or string. Or a string with the Js
    :param profile:

    :return: The htmlObj
    """
    self.css({"cursor": "pointer"})
    self.on("click", jsFncs, profile)
    return self

  @property
  def _js__builder__(self):
    return ''' 
      if(options._children > 0){htmlObj.appendChild(document.createTextNode(data))}
      else{htmlObj.innerHTML = data}
      '''

  def __str__(self):
    return '<span %s>%s</span>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), self.val, self.helper)


class Position(Span):

  def digits(self, flag):
    """
    Description:
    ------------
    Specify if the count should be done from the commas

    Attributes:
    ----------
    :param flag: Boolean (default false)
    """
    self._jsStyles["digits"] = flag
    return self

  def position(self, index, style):
    """
    Description:
    ------------
    Set the CSS format for a specific character at a given position

    Attributes:
    ----------
    :param index: Integer. A number
    :param style: Dictionary. The CSS Style to be used
    """
    self._jsStyles.setdefault("positions", {})[index] = style
    return self

  @property
  def _js__builder__(self):
    return ''' htmlObj.innerHTML = ""; 
        var prevCursor = 0; var content = ""+ data; var shift = 0;
        if (options.digits === true){ shift = content.indexOf(".") + 1; }
        if (typeof options.positions !== 'undefined'){
          const keys = Object.keys(options.positions).sort();
          keys.forEach(function(k){
            var cursor = parseInt(k) + shift;
            var span = document.createElement("span");
            span.innerHTML = content.slice(prevCursor, cursor);
            span.style.display = "inline-block";
            htmlObj.appendChild(span);
            
            var span2 = document.createElement("span");
            span2.innerHTML = content.slice(cursor, cursor+1);
            span2.style.display = "inline-block";
            Object.keys(options.positions[k]).forEach(function(key){
                span2.style[key] = options.positions[k][key]});
            htmlObj.appendChild(span2);
            prevCursor = cursor+1;
          });
          if (content.length > prevCursor){
            var span = document.createElement("span");
            span.innerHTML = content.slice(prevCursor, content.length);
            span.style.display = "inline-block";
            htmlObj.appendChild(span);
          }
        }
        '''

  def __str__(self):
    self._report._props.setdefault('js', {}).setdefault("builders", []).append(self.refresh())
    return '<span %s>%s</span>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), self.val, self.helper)


class Text(Html.Html):
  name, category, callFnc = 'Text', 'Texts', 'text'

  def __init__(self, report, text, color, align, width, height, htmlCode, tooltip, options, helper, profile):
    super(Text, self).__init__(report, text, css_attrs={"color": color, "width": width, "height": height},
                               code=htmlCode, profile=profile)
    self.add_helper(helper)
    self.__options = OptText.OptionsText(self, options)
    self._jsStyles = {"reset": self.options.reset, "markdown": self.options.markdown, "maxlength": self.options.limit_char}
    self.css({'text-align': align})
    if tooltip is not None:
      self.tooltip(tooltip)

  def click(self, jsFncs, profile=False):
    """
    Description:
    ------------
    Add a click event on the text component.
    The style of the mouse on the component will be changed to make the event more visible

    Attributes:
    ----------
    :param jsFncs:
    :param profile:
    """
    self.style.css.cursor = 'pointer'
    return super(Text, self).click(jsFncs, profile)

  @property
  def val(self):
    """
    Description:
    ------------
    Property to get the jquery value of the HTML object in a python HTML object.
    This method can be used in any jsFunction to get the value of a component in the browser.
    This method will only be used on the javascript side, so please do not consider it in your algorithm in Python

    :returns: Javascript string with the function to get the current value of the component
    """
    return self._vals

  @val.setter
  def val(self, val):
    self._vals = val

  @property
  def dom(self):
    """
    Description:
    ------------
    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object

    :rtype: JsHtml.JsHtmlRich
    """
    if self._dom is None:
      self._dom = JsHtml.JsHtmlRich(self, report=self._report)
    return self._dom

  @property
  def options(self):
    """
    Property to set all the possible object for a button

    :rtype: OptText.OptionsText
    """
    return self.__options

  def editable(self):
    """
    Change the componenet properties to be editable if double clicked

    Example
    rptObj.ui.text("This is a text").editable()

    :return: Self to allow the chaining
    """
    self.style.add_classes.text.content_editable()
    self.set_attrs({"contenteditable": "false", "ondblclick": "this.contentEditable=true;this.className='inEdit'",
                    "onblur": "this.contentEditable=false;this.className=''"})
    return self

  @property
  def _js__builder__(self):
    mark_up = self._report.js.string("content", isPyData=False).toStringMarkup()
    return '''
      var content = data;
      if(options.reset){htmlObj.innerHTML = ""}; 
      if(data != ''){ 
        if((options.maxlength != undefined) && (data.length > options.maxlength)){
          content = data.slice(0, options.maxlength); 
          if(options.markdown){htmlObj.innerHTML = %(markUp)s +"..."} else {htmlObj.innerHTML = content +"..."}; 
          htmlObj.title = data} 
        else{
          if(options.markdown){htmlObj.innerHTML = %(markUp)s} else {htmlObj.innerHTML = content}}};
      if(typeof options.css !== 'undefined'){for(var k in options.css){htmlObj.style[k] = options.css[k]}};
      ''' % {"markUp": mark_up}

  def __str__(self):
    if self.options.limit_char and len(self.content) > self.options.limit_char:
      self.set_attrs(name="title", value=self.content)
      if self.options.markdown:
        self._vals = self._report.py.markdown.all(self.content[:self.options.limit_char])
      else:
        self._vals = self.content[:self.options.limit_char]
      return '<div %s>%s...</div>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), self.content, self.helper)

    if self.options.markdown:
      self._vals = self._report.py.markdown.all(self.content)
    return '<div %s>%s</div>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), self.content, self.helper)

  # -----------------------------------------------------------------------------------------
  #                                    EXPORT OPTIONS
  # -----------------------------------------------------------------------------------------
  def to_word(self, document):
    """
    Add to the internal document object a word text section
    """
    document.add_paragraph(self.vals)


class Code(Html.Html):
  name, category, callFnc = 'Code', 'Text', 'code'
  scriptTitle = ''

  def __init__(self, report, vals, color, width, height, htmlCode, options, helper, profile):
    super(Code, self).__init__(report, vals, code=htmlCode, css_attrs={"width": width, "height": height, "color": color}, profile=profile)
    self.add_helper(helper)
    self._jsStyles, self.__editable = options, None
    #self.color = color if color is not None else self._report.theme.greys[9]
    self.css({'display': 'block', 'margin': '5px 0'})

  @property
  def _js__builder__(self):
    return ''' htmlObj.empty();
      if(jsStyles.edit){
        htmlObj.append('<div style="position:relative;float:right;padding:2px 5px 2px 5px;cursor:pointer;background-color:%(blackColor)s;color:%(whiteColor)s">Edit</div>')}
      data.forEach(function(rec){htmlObj.append('<code>'+ rec +'</code><br />')})
      ''' % {"blackColor": self._report.theme.greys[9], "whiteColor": self._report.theme.greys[0]}

  def editable(self, urlPost, title=None):
    self.__editable = urlPost
    self.scriptTitle = title

  def __str__(self):
    if self.__editable is not None:
      self._report.jsOnLoadFnc.add(
        '''
        $('#%(htmlId)s div').on('click', function(event){
          %(jqId)s.empty();
          %(jqId)s.append('<span style="position:relative;float:right;padding:2px 5px 2px 5px;cursor:pointer;background-color:%(backGroundColor)s;color:%(whiteColor)s">Save</span>');
          %(jqId)s.attr('contenteditable','true'); %(jqId)s.css('padding', '0 0 0 5px');
          recordSet_%(htmlId)s.forEach(function(rec) { %(jqId)s.append(rec + "\\n")});
          %(jqId)s.append( "<br />");
          $('#%(htmlId)s span').on('click', function(event){
            var content = %(jqId)s.text();
            $.post("/%(url)s", {content: content.slice(4, content.length), title: '%(title)s'}, function(data){location.reload()})});   
          })''' % {'jqId': self.jqId, "url": self.editable, 'htmlId': self.htmlId, 'backGroundColor': self._report.theme.colors[5],
                   'whiteColor': self._report.theme.greys[0], 'title': self.scriptTitle})
    return '<div %s></div>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), self.helper)


class Pre(Html.Html):
  name, category, callFnc = 'Pre formatted text', 'Texts', 'preformat'

  def __init__(self, report, vals, color, width, height, htmlCode, dataSrc, options, helper, profile):
    super(Pre, self).__init__(report, vals, code=htmlCode, css_attrs={"width": width, 'height': height, 'color': color},
                              profile=profile, dataSrc=dataSrc)
    self._jsStyles = options
    self.css({"text-align": 'left'})
    self.add_helper(helper)

  @property
  def dom(self):
    """
    Javascript Functions

    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object

    :rtype: JsHtml.JsHtmlRich
    """
    if self._dom is None:
      self._dom = JsHtml.JsHtmlRich(self, report=self._report)
    return self._dom

  def selectable(self, flag=False):
    """
    Description:
    ------------
    Make the label component not selectable.

    This will be done by adding the class CssTextNotSelectable to the component

    Attributes:
    ----------
    :param flag: Boolean.

    :return: self to allow the chains
    """
    if not flag:
      self.style.add_classes.text.no_selection()
    return self

  @property
  def _js__builder__(self):
    markdown = self._report.js.string("data", isPyData=False).toStringMarkup()
    return '''if(options.markdown){htmlObj.innerHTML = %(markdown)s} else{htmlObj.innerHTML = data}''' % {"markdown": markdown}

  def __str__(self):
    return '<pre %s>%s</pre>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), self.val, self.helper)


class Paragraph(Html.Html):

  name, category, callFnc = 'Paragraph', 'Texts', 'paragraph'

  def __init__(self, report, text, color, background_color, border, width, height, htmlCode, encoding, dataSrc,
               helper, profile):
    jsStyles, tmpText = {"reset": True, 'markdown': True, "classes": []}, []
    if not isinstance(text, list):
      content = []
      for line in text.strip().split("\n"):
        content.append(line.strip())
      text = [" ".join(content)]
    for t in text:
      cssStyles = re.search(" css\{(.*)\}", t)
      jsAttr = {}
      if cssStyles is not None:
        content = t.replace(cssStyles.group(0), '').decode(encoding) if hasattr(t, 'decode') else t.replace(cssStyles.group(0), '')
        tmpText.append(content)
        for cssAttr in cssStyles.group(1).split(","):
          cssKey, cssVal = cssAttr.split(":")
          jsAttr[cssKey.strip()] = cssVal.strip()
      else:
        if hasattr(t, 'decode'):
          tmpText.append(t.decode(encoding))
        else:
          tmpText.append(t)
      jsStyles["classes"].append(jsAttr)
    super(Paragraph, self).__init__(report, tmpText, code=htmlCode, css_attrs={'color': color, "width": width,
           "height": height, "background-color": background_color}, dataSrc=dataSrc, profile=profile)
    self.add_helper(helper)
    self._jsStyles = jsStyles
    if border:
      self.css('border', '1px solid %s' % self._report.theme.greys[9])
    self.css({'text-align': 'justify', 'margin-top': '3px', "text-justify": 'distribute'})

  @property
  def _js__builder__(self):
    markdown = self._report.js.string("line", isPyData=False).toStringMarkup()
    return '''
      if (typeof options.reset === 'undefined' || options.reset){htmlObj.innerHTML = ''};
      if (typeof data === 'string' || data instanceof String){data = data.split('\\n')}; 
      data.forEach(function(line, i){
        var p = document.createElement('p'); p.innerHTML = %(markdown)s;
        htmlObj.appendChild(p)});
      if(typeof options.css !== 'undefined'){for(var k in options.css){htmlObj.style[k] = options.css[k]}}
      ''' % {"markdown": markdown}

  @property
  def dom(self):
    """
    Javascript Functions

    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object

    :rtype: JsHtml.JsHtmlRich
    """
    if self._dom is None:
      self._dom = JsHtml.JsHtmlRich(self, report=self._report)
    return self._dom

  def __str__(self):
    self._report._props.setdefault('js', {}).setdefault("builders", []).append(self.refresh())
    return '<div %s></div>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), self.helper)

  # -----------------------------------------------------------------------------------------
  #                                    EXPORT OPTIONS
  # -----------------------------------------------------------------------------------------
  def to_word(self, document):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = document.add_paragraph(self.vals)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

  def to_xls(self, workbook, worksheet, cursor):
    """

    :param workbook:
    :param worksheet:
    :param cursor:

    :link xlxWritter Documentation: https://xlsxwriter.readthedocs.io/format.html
    """
    worksheet.write(cursor['row'], 0, "\n".join(self.vals))
    cursor['row'] += 2


class BlockQuote(Html.Html):
  name, category, callFnc = 'Block quotation', 'Texts', 'blockquote'

  def __init__(self, report, text, author, color, width, height, htmlCode, helper, profile):
    super(BlockQuote, self).__init__(report, {'text': text, 'author': author}, code=htmlCode, profile=profile,
                                     css_attrs={"width": width, "height": height, 'color': color, "white-space": 'nowrap'})
    self.add_helper(helper)
    self._jsStyles = {"reset": True, 'markdown': True}

  @property
  def _js__builder__(self):
      return '''var div = htmlObj.querySelector('div'); div.innerHTML = ''; console.log(data);
          data.text.split("\\n").forEach(function(rec) {
            var p = document.createElement("p");
            p.style.margin = 0; p.style.padding = 0; p.innerHTML = rec; div.appendChild(p) });
          if(data.author != null){htmlObj.querySelector('div:last-child').innerHTML = '<small>by '+ data.author +'<cite></cite></small>'}'''

  def __str__(self):
    self._report._props.setdefault('js', {}).setdefault("builders", []).append(self.refresh())
    return '''
      <blockquote %s>
          <div style="padding:5px;border-left:4px solid %s"></div>
          <div style="text-align:right"></div>
      </blockquote>%s''' % (self.get_attrs(pyClassNames=self.style.get_classes()), self._report.theme.colors[9], self.helper)

  # -----------------------------------------------------------------------------------------
  #                                    MARKDOWN SECTION
  # -----------------------------------------------------------------------------------------
  @staticmethod
  def matchMarkDownBlock(data): return re.match(">>>Quote:(.*)", data[0])

  @staticmethod
  def matchEndBlock(data): return data.endswith("<<<")

  @classmethod
  def convertMarkDownBlock(cls, data, report):
    author = data[0].split(':')[-1]
    getattr(report, 'blockquote')("\n".join([val for val in data[1:-1]]), author)
    return ["report.blockquote(%s, '%s')" % (json.dumps([val for val in data[1:-1]]), author)]

  @classmethod
  def jsMarkDown(self, vals): return [">>>Quote:%s" % self.val['author'], [rec for rec in vals['text']], "<<<"]


class Title(Html.Html):
  name, category, callFnc = 'Title', 'texts', 'title'

  def __init__(self, report, text, level, name, contents, color, picture, icon, marginTop, htmlCode, width,
               height, align, options, profile):

    self.__options = OptText.OptionsTitle(self, options)
    cssStyles, jsStyles = re.search(" css\{(.*)\}", text), options
    if cssStyles is not None:
      text = text.replace(cssStyles.group(0), '')
      for cssAttr in cssStyles.group(1).split(","):
        cssKey, cssVal = cssAttr.split(":")
        jsStyles[cssKey.strip()] = cssVal.strip()
    super(Title, self).__init__(report, text, code=htmlCode, css_attrs={"width": width, "height": height}, profile=profile)
    self._jsStyles = jsStyles
    self._name, self.level, self.picture = name, level, picture
    self.add_icon(icon)
    if contents is not None:
      self._name = contents.add(text, level or 1, name)
    if level is not None:
      getattr(self.style.add_classes.text, "title_%s" % level)()
      self.css({'color': color, 'margin': '%spx 0 5px 0' % marginTop})
    else:
      self.style.add_classes.text.title()
      self.css({'margin': '%spx 0 5px 0' % marginTop})
    if align == 'center':
      self.css({'margin': '5px auto 10px auto', 'display': 'block', 'text-align': 'center'})
    elif align is not None:
      self.css({'margin': '5px auto 10px auto', 'display': 'block', 'text-align': align})
    else:
      self.css({'display': 'block', "margin-right": "10px"})
    if hasattr(report, '_content_table') and self.__options.content_table:
      # Special attribute set in the base component interface
      div = self._report.ui.div(htmlCode="%s_anchor" % self.htmlId)
      if self._report.body.css('padding-top') is None:
        div.style.css.margin_top = - 10
      else:
        div.style.css.margin_top = - int(self._report.body.css('padding-top')[:-2]) - 10
      div.style.css.position = "absolute"
      div.style.css.z_index = -1
      report._content_table.add(text, level or 4, "#%s_anchor" % self.htmlId)
      report._content_table[-1].click([
        self.dom.transition(["color", "font-size"], ['red', '102%'], duration=[0.5, 0.5], reverse=True)])

  @property
  def style(self):
    """
    Description:
    ------------

    :rtype: GrpCls.ClassHtmlEmpty
    """
    if self._styleObj is None:
      self._styleObj = GrpCls.ClassHtmlEmpty(self)
    return self._styleObj

  @property
  def dom(self):
    """
    Description:
    ------------
    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object

    :rtype: JsHtml.JsHtmlRich
    """
    if self._dom is None:
      self._dom = JsHtml.JsHtmlRich(self, report=self._report)
    return self._dom

  @property
  def options(self):
    """
    Property to set all the possible object for a button

    :rtype: OptText.OptionsText
    """
    return self.__options

  @property
  def _js__builder__(self):
    markdown = self._report.js.string("data", isPyData=False).toStringMarkup()
    return '''
      if(options.markdown){htmlObj.innerHTML = %(markdown)s} else{htmlObj.innerHTML = data};
      if(typeof options.css !== 'undefined'){for(var k in options.css){htmlObj.style[k] = options.css[k]}}''' % {"markdown": markdown}

  def __str__(self):
    anchor_name = ' name="%s"' % self._name if self._name is not None else ''
    if self.picture is not None:
      path = "/img/%s" % self._report.run.report_name
      # Check if the file is available in the default directory
      # Otherwise raise an exception
      filePath = os.path.join(self._report.run.local_path, "static", self.picture)
      if not os.path.exists(filePath):
        raise Exception("Missing file %s in %s" % (self.picture, os.path.join(self._report.run.local_path, "static")))

      return '<div %s><img src="%s/%s" />&nbsp;<a%s></a>%s%s</div>' % (self.get_attrs(pyClassNames=self.style.get_classes()), path, self.picture, anchor_name, self.val, self.helper)

    return '<div %s><a%s></a>%s%s</div>' % (self.get_attrs(pyClassNames=self.style.get_classes()), anchor_name, self.val, self.helper)

  # -----------------------------------------------------------------------------------------
  #                                    MARKDOWN SECTION
  # -----------------------------------------------------------------------------------------
  @staticmethod
  def matchMarkDown(val): return val.startswith("# ") or val.startswith("## ") or val.startswith("### ") or val.startswith("#### ")

  @classmethod
  def convertMarkDown(cls, val, regExpResult, report=None):
    helper, infoData = re.search("epyk:info<(.*)>", val), None
    if helper is not None:
      infoData = helper.group(1)
      val = val.replace("epyk:info<%s>" % infoData, "")
    for i in range(1, 5):
      startIndex = i+1
      if val.startswith("%s " % "".join(["#"] * i)):
        if report is not None:
          title = getattr(report.ui.texts, cls.callFnc)(val[startIndex:], level=i)
          if infoData is not None:
            title.info(infoData)
        return ["report.ui.%s.title('%s', level=%s)" % (cls.category, val[startIndex:], i)]

    if report is not None:
      getattr(report.ui.texts, cls.callFnc)(val[2:])
    return ["report.ui.%s.%s('%s')" % (val[2:], cls.category, cls.callFnc)]

  @classmethod
  def jsMarkDown(cls, vals, level=None):
    if level == 1: return "# %s" % vals
    if level == 2: return "## %s" % vals
    if level == 3: return "### %s" % vals
    return "#### %s" % vals

  # -----------------------------------------------------------------------------------------
  #                                    EXPORT OPTIONS
  # -----------------------------------------------------------------------------------------
  def to_ppt(self, document, currentSlide, newSlide=False):
    if newSlide:
      blankSlide = document.slide_layouts[6]
      currentSlide = document.slides.add_slide(blankSlide)
    textbox = currentSlide.shapes.add_textbox(0, 0, 100, 100)
    textbox.text_frame.text = self.vals
    return currentSlide

  def to_word(self, document):
    if self.level is not None:
      document.add_heading(self.vals, self.level)
    else:
      document.add_heading(self.vals)

  def to_xls(self, workbook, worksheet, cursor):
    """

    Documentation
    https://xlsxwriter.readthedocs.io/format.html

    :param workbook:
    :param worksheet:
    :param cursor:

    :return:
    """
    cell_format = workbook.add_format({'bold': True, 'font_color': self.color, 'font_size': self.size})
    worksheet.write(cursor['row'], 0, self.vals, cell_format)
    cursor['row'] += 2


class Numeric(Html.Html):
  name, category, callFnc = 'Number', 'Number', 'number'
  __reqJs = ['accounting']

  def __init__(self, report, number, title, label, icon, color, tooltip, htmlCode, options, helper, width, profile):
    super(Numeric, self).__init__(report, number, code=htmlCode, profile=profile, css_attrs={"width": width, "color": color})
    # Add the components label and icon
    self.add_label(label, css={"float": None, "width": 'none'})
    self.add_icon(icon)
    self.add_helper(helper, css={"line-height": '20px'})
    self.add_title(title, level=4, css={"margin-bottom": 0}, options={'content_table': False})

    # Update the CSS Style of the component
    self.css({'text-align': 'center', 'display': 'inline-block'})
    self.tooltip(tooltip)
    self.__options = OptText.OptionsNumber(self, options)

  def money(self, symbol="", digit=0, thousand_sep=".", decimal_sep=",", format="%s%v"):
    """
    Description:
    -----------
    Format any number into currency

    Related Pages:
    --------------
    http://openexchangerates.github.io/accounting.js/

    Attributes:
    ----------
    :param symbol: String. custom symbol
    :param digit: Integer. Number of digit
    :param thousand_sep: String. The thousand separator
    :param decimal_sep: String. The decimal separator
    """
    self.options.symbol = symbol
    self.options.format = format
    self.options.digit = digit
    self.options.thousand_sep = thousand_sep
    self.options.decimal_sep = decimal_sep
    return self

  def number(self, digits=0, thousand_sep=',', decimal_sep=","):
    """
    Description:
    -----------
    Format a number with custom precision and localisation

    Related Pages:
    --------------
    http://openexchangerates.github.io/accounting.js/

    Attributes:
    ----------
    :param digits: Integer. Number of digit
    :param thousand_sep: String. The thousand separator
    """
    self._jsStyles["type_number"] = "number"
    self.options.digits = digits
    self.options.thousand_sep = thousand_sep
    self.options.decimal_sep = decimal_sep
    return self

  @property
  def options(self):
    """
    Property to set all the possible object for a button

    :rtype: OptText.OptionsNumber
    """
    return self.__options

  @property
  def _js__builder__(self):
    return '''
      if (options.type_number == 'money'){ htmlObj.querySelector('font').innerHTML = accounting.formatMoney(data, options.symbol, options.digits, options.thousand_sep, options.decimal_sep, options.format) }
      else { htmlObj.querySelector('font').innerHTML = accounting.formatNumber(data, options.digits, options.thousand_sep, options.decimal_sep) }      
      '''

  def __str__(self):
    self._report._props.setdefault('js', {}).setdefault("builders", []).append(self.refresh())
    return "<div %s><font style='vertical-align:middle;height:100%%;padding:0;margin:0;display:inline-block'>%s</font>%s</div>" % (self.get_attrs(pyClassNames=self.style.get_classes()), self.val, self.helper)


class Highlights(Html.Html):
  name, category, callFnc = 'Highlights', 'Texts', 'highlights'
  __reqCss, __reqJs = ['bootstrap'], ['bootstrap']

  def __init__(self, report, text, title, icon, type, color, width, height, htmlCode, helper, profile):
    super(Highlights, self).__init__(report, text, css_attrs={"width": width, "height": height}, code=htmlCode, profile=profile)
    self.add_helper(helper)
    self.color = color if color is not None else self._report.theme.greys[9]
    # Add the components title and icon
    self.add_title(title, css={"width": "none", "font-weight": 'bold'}, options={'content_table': False})
    self.add_icon(icon, {"float": "left"})
    # Change the style of the component
    self.css({"margin": "5px", 'padding': "5px"})
    self.attr['class'].add('alert alert-%s' % type)
    self.set_attrs(name='role', value="alert")

  def __str__(self):
    return "<div %s><div>%s</div></div>%s" % (self.get_attrs(pyClassNames=self.style.get_classes()), self.val, self.helper)


class Fieldset(Html.Html):
  name, category, callFnc = 'Fieldset', 'Texts', 'fieldset'

  def __init__(self, report, legend, width, height, helper, profile):
    super(Fieldset, self).__init__(report, legend, css_attrs={"width": width, "height": height}, profile=profile)
    self.add_helper(helper)
    self.css({'padding': '5px', 'border': '1px groove %s' % self._report.theme.greys[3], 'display': 'block', 'margin': '5px 0'})
    self.components = []

  @property
  def _js__builder__(self):
    return '''htmlObj.firstChild.innerHTML = data; 
      if(typeof options.css !== 'undefined'){Object.keys(options.css).forEach(function(key){htmlObj.firstChild.style[key] = options.css[key]})}'''

  def __add__(self, htmlObj):
    """ Add items to a container """
    htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    self.components.append(htmlObj)
    return self

  def __getitem__(self, i):
    return self.components[i]

  def __str__(self):
    str_div = "".join([v.html() if hasattr(v, 'html') else str(v) for v in self.components])
    return '<fieldset %s><legend style="width:auto">%s</legend>%s</fieldset>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), self.val, str_div, self.helper)
